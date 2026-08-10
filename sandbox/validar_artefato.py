"""Validação dos artefatos que o agente entrega (.xlsx/.xlsm, .docx, .pdf).

Este módulo roda DENTRO do sandbox: `agent/outputs.js` lê este arquivo, anexa
uma linha de driver com a lista de arquivos e manda o texto inteiro para o
`run_python`. Por isso ele não importa nada do projeto e só depende do que o
sandbox tem instalado (openpyxl, python-docx, pypdf).

Ele viveu 233 linhas dentro de uma template string de JavaScript, onde era
**impossível de testar**: o F-23 da auditoria era exatamente isso — o filtro que
escolhe os arquivos tinha teste, e o código que decide "este arquivo está bom"
não tinha nenhum. Um validador sem teste é pior que nenhum, porque a entrega
passa a dizer "verificado" apoiada em algo que ninguém conferiu.

O par deste arquivo é `validar_artefato_test.py`, que gera .xlsx/.docx/.pdf de
verdade e cobra cada veredito.
"""
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile

ERRS = ("#REF!", "#VALUE!", "#DIV/0!", "#N/A", "#NAME?", "#NULL!", "#NUM!", "#SPILL!", "#CALC!")


class Config:
    """Limites da validação.

    Eram constantes de módulo lidas do ambiente no import, e o orçamento de
    recálculo era um global mutável — o que fazia um teste contaminar o
    seguinte. Aqui viram estado por chamada. No sandbox o efeito é o mesmo (o
    ambiente não muda no meio do processo); no teste, dá para exercitar o teto
    de células sem gerar uma planilha de 40 mil células.
    """

    def __init__(self, max_cells=40000, recalc_enabled=True, recalc_timeout=25, recalc_max_files=2):
        self.max_cells = max_cells
        self.recalc_enabled = recalc_enabled
        self.recalc_timeout = recalc_timeout
        # Lista de um elemento porque `recalc` desconta o orçamento por dentro.
        self.budget = [recalc_max_files]

    @classmethod
    def from_env(cls, env=None):
        env = os.environ if env is None else env
        return cls(
            max_cells=int(env.get("VALIDATE_MAX_CELLS", "40000")),
            recalc_enabled=env.get("VALIDATE_RECALC", "true").lower() == "true",
            recalc_timeout=int(env.get("SANDBOX_RECALC_TIMEOUT_S", "25")),
            recalc_max_files=int(env.get("RECALC_MAX_FILES", "2")),
        )

def scan_errors(wb, cap):
    errs = scanned = 0
    capped = False
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if scanned >= cap:
                    capped = True
                    break
                v = c.value
                if v is not None:
                    scanned += 1
                if isinstance(v, str) and any(e in v for e in ERRS):
                    errs += 1
            if capped:
                break
        if capped:
            break
    return errs, capped

def recalc(path, cfg):
    tmp = tempfile.mkdtemp(prefix="lo_recalc_")
    prof = os.path.join(tmp, "profile", "user")
    os.makedirs(prof, exist_ok=True)
    with open(os.path.join(prof, "registrymodifications.xcu"), "w", encoding="utf-8") as f:
        f.write('<?xml version="1.0" encoding="UTF-8"?>'
                '<oor:items xmlns:oor="http://openoffice.org/2001/registry" '
                'xmlns:xs="http://www.w3.org/2001/XMLSchema" '
                'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
                '<item oor:path="/org.openoffice.Office.Calc/Formula/Load">'
                '<prop oor:name="OOXMLRecalcMode" oor:op="fuse"><value>0</value></prop></item>'
                '<item oor:path="/org.openoffice.Office.Calc/Formula/Load">'
                '<prop oor:name="ODFRecalcMode" oor:op="fuse"><value>0</value></prop></item>'
                '</oor:items>')
    outdir = os.path.join(tmp, "out")
    os.makedirs(outdir, exist_ok=True)
    try:
        subprocess.run(["soffice", "-env:UserInstallation=file://" + os.path.join(tmp, "profile"),
                        "--headless", "--calc", "--convert-to", "xlsx:Calc MS Excel 2007 XML",
                        "--outdir", outdir, path],
                       timeout=cfg.recalc_timeout, capture_output=True, check=False)
    except Exception:
        return None, tmp
    cand = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".xlsx")
    return (cand if os.path.exists(cand) else None), tmp

def recalc_took(orig, calc):
    keys = set()
    for ws in orig.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if isinstance(c.value, str) and c.value.startswith("="):
                    keys.add((ws.title, c.coordinate))
                    if len(keys) > 40:
                        break
            if len(keys) > 40:
                break
        if len(keys) > 40:
            break
    if not keys:
        return True
    for ws in calc.worksheets:
        for row in ws.iter_rows():
            for c in row:
                # Qualquer valor em cache (número OU string de erro tipo #DIV/0!)
                # prova que o recálculo aconteceu. Antes exigia só numérico, o
                # que descartava o recálculo quando TODAS as fórmulas viravam
                # erro — exatamente o caso que mais precisamos detectar.
                if (ws.title, c.coordinate) in keys and c.value is not None:
                    return True
    return False

def _col_to_num(col):
    n = 0
    for ch in col.upper():
        n = n * 26 + (ord(ch) - 64)
    return n

def _parse_ref(ref):
    m = re.match(r"^(?:'([^']+)'|([^!]+))!(.+)$", ref.strip())
    if not m:
        return None
    sheet = m.group(1) or m.group(2)
    rng = m.group(3).replace("$", "")
    cells = rng.split(":")
    def cell(c):
        mm = re.match(r"^([A-Za-z]+)(\d+)$", c)
        return (_col_to_num(mm.group(1)), int(mm.group(2))) if mm else None
    a = cell(cells[0])
    b = cell(cells[-1]) if len(cells) > 1 else a
    if not a or not b:
        return (sheet, None)
    return (sheet, (a[0], a[1], b[0], b[1]))

def _range_has_number(wb, sheet, coords):
    # True se houver ao menos uma célula NUMÉRICA no intervalo. Usado para pegar
    # gráfico cuja série de VALORES aponta para um intervalo sem números (ex.:
    # coluna "Total" declarada no cabeçalho mas deixada vazia pelo modelo).
    try:
        ws = wb[sheet]
    except Exception:
        return True  # na dúvida, não acusa
    c1, r1, c2, r2 = coords
    for r in range(r1, r2 + 1):
        for c in range(c1, c2 + 1):
            v = ws.cell(row=r, column=c).value
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                return True
            if isinstance(v, str) and v.startswith("="):
                return True  # fórmula: assume que produz número
    return False

def check_charts(p, wb):
    # Valida os GRÁFICOS do .xlsx (openpyxl descarta charts ao carregar, então
    # lemos o XML direto do zip). Modelos geram gráficos com referências
    # quebradas — intervalos INVERTIDOS (ex.: C2:B2), aba inexistente, sem série
    # ou uma série de VALORES apontando para um intervalo vazio (coluna sem
    # dados) — que abririam vazios/errados no Excel, e nada verificava isso.
    sheetnames = wb.sheetnames
    problems = []
    n = 0
    try:
        z = zipfile.ZipFile(p)
    except Exception:
        return 0, problems
    for cn in [x for x in z.namelist() if re.match(r"xl/charts/chart\d+\.xml", x)]:
        n += 1
        try:
            xml = z.read(cn).decode("utf-8", "replace")
        except Exception:
            continue
        refs = re.findall(r"<(?:\w+:)?f>([^<]+)</(?:\w+:)?f>", xml)
        sers = len(re.findall(r"<(?:\w+:)?ser>", xml))
        # referências que são SÉRIE DE VALORES (dentro de <val>...</val>) — só
        # essas precisam ter números; categorias podem ser texto.
        val_refs = []
        for blk in re.findall(r"<(?:\w+:)?val>(.*?)</(?:\w+:)?val>", xml, re.S):
            val_refs += re.findall(r"<(?:\w+:)?f>([^<]+)</(?:\w+:)?f>", blk)
        if not refs:
            problems.append("grafico sem referencias de dados")
            continue
        if sers == 0:
            problems.append("grafico sem series")
        for ref in refs:
            parsed = _parse_ref(ref)
            if not parsed:
                continue
            sheet, coords = parsed
            if sheet not in sheetnames:
                problems.append("grafico referencia aba inexistente '" + sheet + "'")
                continue
            if coords:
                c1, r1, c2, r2 = coords
                if c1 > c2 or r1 > r2:
                    problems.append("grafico com intervalo invertido/degenerado (" + ref + ")")
                    continue
                if ref in val_refs and not _range_has_number(wb, sheet, coords):
                    problems.append("grafico com serie de valores vazia (" + ref + ")")
    return n, problems

def check_xlsx(p, cfg=None):
    cfg = cfg or Config.from_env()
    from openpyxl import load_workbook
    wb = load_workbook(p)
    sheets = len(wb.sheetnames)
    errs_lint, capped = scan_errors(wb, cfg.max_cells)
    errs_calc = 0
    did_recalc = False
    if cfg.recalc_enabled and cfg.budget[0] > 0:
        cfg.budget[0] -= 1
        calc_path, tmp = recalc(p, cfg)
        try:
            if calc_path:
                calc_wb = load_workbook(calc_path, data_only=True)
                if recalc_took(wb, calc_wb):
                    did_recalc = True
                    errs_calc, _ = scan_errors(calc_wb, cfg.max_cells)
        except Exception:
            pass
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    total = errs_lint + errs_calc
    chart_n, chart_problems = check_charts(p, wb)
    parts = [str(sheets) + " abas"]
    if total:
        parts.append(str(total) + " celula(s) com erro de formula")
    elif did_recalc:
        parts.append("formulas recalculadas: sem erros")
    else:
        parts.append("formulas NAO recalculadas (verificacao parcial)")
    if chart_problems:
        parts.append(str(len(chart_problems)) + " problema(s) em grafico: " + chart_problems[0])
    elif chart_n:
        parts.append(str(chart_n) + " grafico(s) ok")
    if capped:
        parts.append("varredura limitada a " + str(cfg.max_cells) + " celulas")
    return {"ok": total == 0 and not chart_problems, "info": "; ".join(parts)}

def check_docx(p):
    from docx import Document
    d = Document(p)
    n_par = len([x for x in d.paragraphs if x.text.strip()])
    empty = (n_par == 0 and len(d.tables) == 0 and len(d.inline_shapes) == 0)
    return {"ok": not empty, "info": str(len(d.paragraphs)) + " paragrafos, " + str(len(d.tables)) + " tabelas" + (" - documento vazio" if empty else "")}

def check_pdf(p):
    from pypdf import PdfReader

    n = len(PdfReader(p).pages)
    return {"ok": n > 0, "info": str(n) + " paginas"}


def validar(files, base="/workspace", cfg=None):
    """Roteia cada arquivo para o validador da sua extensão.

    Extensão desconhecida sai `ok: True` com `info` vazio de propósito: o
    validador não sabe julgar aquele formato, e reprovar o que não se sabe ler
    transformaria "não verifiquei" em "está errado".

    Arquivo que não abre é `ok: False` — aí a falha é sobre o arquivo, não
    sobre a nossa ignorância. A mensagem vai truncada em 120 caracteres porque
    stack trace de biblioteca não cabe (nem ajuda) na tela do usuário.
    """
    cfg = cfg or Config.from_env()
    out = []
    for rel in files:
        p = os.path.join(base, rel)
        r = {"path": rel, "ok": True, "info": ""}
        try:
            ext = rel.lower().rsplit(".", 1)[-1]
            if ext in ("xlsx", "xlsm"):
                r.update(check_xlsx(p, cfg))
            elif ext == "pdf":
                r.update(check_pdf(p))
            elif ext == "docx":
                r.update(check_docx(p))
        except Exception as e:
            r["ok"] = False
            r["info"] = ("nao abre: " + str(e))[:120]
        out.append(r)
    return out
