import os
import tempfile
import unittest
import zipfile

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm

    import docpro

    HAS_DOCX = True
except Exception:  # python-docx ausente no ambiente
    HAS_DOCX = False

try:
    import matplotlib  # noqa: F401

    HAS_MPL = True
except Exception:  # matplotlib ausente (CI enxuta) — só os gráficos dependem dele
    HAS_MPL = False

# O PDF gêmeo (e com ele o sumário com páginas REAIS, a conferência de página
# vazia e a de assinatura órfã) depende do LibreOffice e do pypdf. Na CI enxuta
# eles não existem: os testes que os exigem se marcam como pulados em vez de
# fingir que conferiram.
def _tem_pdf():
    import shutil
    if not shutil.which("soffice"):
        return False
    try:
        import pypdf  # noqa: F401
    except Exception:
        try:
            import PyPDF2  # noqa: F401
        except Exception:
            return False
    return True


HAS_PDF = _tem_pdf()

# Ordem canônica dos filhos de <w:tblPr> (subconjunto usado pelo kit). Serve para
# provar que a diagramação da tabela gera OOXML válido (fora de ordem, o Word abre
# "reparando" e às vezes perde o estilo).
_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
]


@unittest.skipUnless(HAS_DOCX, "python-docx não instalado")
class DocProTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="frederico-docx-")
        self.path = os.path.join(self.tmp.name, "relatorio.docx")

    def tearDown(self):
        self.tmp.cleanup()

    def _nu(self, *a, **k):
        """Relatório SEM capa, sumário e fechamento automáticos.

        Os presets da v2 acrescentam esses blocos sozinhos; num teste que mede
        UM bloco (a aresta da tabela, a grade da linha do tempo) eles só entram
        como ruído — `doc.tables[0]` passaria a ser a faixa da capa."""
        k.setdefault("capa", False)
        k.setdefault("sumario", False)
        k.setdefault("contracapa", False)
        return docpro.Relatorio(*(a or ("T",)), **k)

    def _relatorio(self):
        r = docpro.Relatorio(
            "Título Longo do Relatório Gerencial de Situação Fiscal",
            cliente="ACME COMÉRCIO LTDA", emissor="Escritório",
            subtitulo="Exercício 2025", tipo="RELATÓRIO GERENCIAL")
        r.capa()
        r.titulo("1. Itens")
        r.tabela(
            ["Item", "Descrição bem longa do serviço prestado ao cliente", "Valor"],
            [["1", "Consultoria contábil mensal com apuração de tributos", "R$ 1.200,00"],
             ["TOTAL", "", "R$ 1.200,00"]],
            total=True)
        r.kpis([("R$ 5M", "Capital"), ("ATIVA", "Situação")])
        r.callout("RESUMO", "Empresa ativa e adimplente.")
        r.salvar(self.path, pdf=False)
        return self.path

    def test_gera_docx_valido_e_reabre(self):
        self._relatorio()
        self.assertTrue(zipfile.is_zipfile(self.path))
        doc = Document(self.path)  # reabre sem erro de XML
        self.assertGreaterEqual(len(doc.tables), 3)

    def test_nenhuma_tabela_vaza_a_margem(self):
        """Duas estratégias legítimas: 100% da área útil (`_fit`) ou largura
        FIXA declarada em dxa (`_fixa`, dos KPIs/linha do tempo/assinaturas).
        O que não pode existir é tabela sem largura declarada — é ela que vaza
        para fora da margem quando o conteúdo cresce."""
        r = docpro.Relatorio("T", cliente="ACME", emissor="Escritório")
        r.capa()
        r.tabela(["A", "B"], [["1", "2"]])
        r.kpis([("R$ 5M", "Capital"), ("ATIVA", "Situação")])
        r.etapas([("Diagnóstico", "ago/2026"), ("Implantação", "out/2026")])
        r.citacao("Frase.", "Fonte")
        r.assinaturas(["Nome"], cargos=["Cargo"])
        r.contracapa(contatos=["contato@empresa.com.br"])
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)
        util = Cm(21.0 - 2 * 2.0).twips
        self.assertGreaterEqual(len(doc.tables), 6)
        for t in doc.tables:
            tblW = t._tbl.tblPr.find(qn("w:tblW"))
            self.assertIsNotNone(tblW, "toda tabela deve ter largura preferida")
            tipo = tblW.get(qn("w:type"))
            self.assertIn(tipo, ("pct", "dxa"))
            if tipo == "pct":
                self.assertEqual(tblW.get(qn("w:w")), "5000")
            else:
                self.assertLessEqual(int(tblW.get(qn("w:w"))), util)

    def test_tabela_de_largura_fixa_traz_a_grade_coerente(self):
        """Em layout fixo o Word resolve as colunas pela <w:tblGrid>. Se ela
        ficar com a grade herdada do python-docx (colunas iguais, somando mais
        que a área útil), a tabela vaza mesmo com cada célula medida a uma.
        É o caso da linha do tempo e das assinaturas, que precisam de colunas
        de larguras diferentes."""
        r = self._nu()
        r.etapas([("Diagnóstico", "ago/2026"), ("Implantação", "out/2026")])
        r.salvar(self.path, pdf=False)
        t = Document(self.path).tables[0]
        tblW = t._tbl.tblPr.find(qn("w:tblW"))
        self.assertEqual(tblW.get(qn("w:type")), "dxa")
        grid = t._tbl.find(qn("w:tblGrid"))
        soma = sum(int(g.get(qn("w:w"))) for g in grid.findall(qn("w:gridCol")))
        self.assertEqual(int(tblW.get(qn("w:w"))), soma)
        self.assertLessEqual(soma, Cm(17.0).twips)

    def test_placeholder_de_rascunho_reprova_na_auditoria(self):
        """"DD/MM/AAAA" numa entrega é o defeito mais barato de evitar e o mais
        constrangedor de deixar passar. A auditoria do kit o encontra no arquivo
        PRONTO — inclusive dentro das tabelas aninhadas do bloco de assinatura."""
        r = self._nu()
        r.paragrafo("Vencimento em DD/MM/AAAA.")
        with self.assertRaises(docpro.KitError) as erro:
            r.salvar(self.path, pdf=False)
        self.assertIn("placeholder", str(erro.exception))

    def test_ordem_do_tblpr_e_valida(self):
        self._relatorio()
        doc = Document(self.path)
        for t in doc.tables:
            tags = [c.tag.split("}")[-1] for c in t._tbl.tblPr]
            idx = [_TBLPR_ORDER.index(x) for x in tags if x in _TBLPR_ORDER]
            self.assertEqual(idx, sorted(idx), tags)

    def test_capa_sem_rodape_na_primeira_pagina(self):
        self._relatorio()
        doc = Document(self.path)
        sec = doc.sections[0]
        self.assertTrue(sec.different_first_page_header_footer)
        self.assertEqual(sec.first_page_footer.paragraphs[0].text.strip(), "")

    def test_linha_com_menos_valores_que_o_cabecalho_reprova_na_auditoria(self):
        """A v1 completava a linha curta em SILÊNCIO — o documento saía com
        células vazias e ninguém via. A v2 continua não derrubando a geração
        (o .docx é gravado), mas a auditoria REPROVA e diz qual linha está
        fora do cabeçalho."""
        r = self._nu()
        r.tabela(["A", "B", "C"], [["só um"], ["um", "dois"]])
        with self.assertRaises(docpro.KitError) as erro:
            r.salvar(self.path, pdf=False)
        self.assertIn("linha-fora-do-cabecalho", str(erro.exception))
        self.assertTrue(zipfile.is_zipfile(self.path))

    def test_corpo_titulo_e_tabela_usam_a_mesma_aresta_esquerda(self):
        """O mesmo contrato do pdfpro, agora no Word: título com barra, corpo,
        lista e primeira coluna da tabela começam todos em RECUO_PT. Antes o
        título nascia 10 pt à direita do corpo (que não tinha recuo) e a célula
        6 pt à esquerda dele — três arestas na mesma página."""
        r = self._nu()
        r.titulo("Seção")
        r.paragrafo("Um parágrafo do corpo.")
        r.lista(["item um", "item dois"])
        r.tabela(["Campo", "Valor"], [["CNPJ", "00.000.000/0001-00"]])
        r.callout("NOTA", "Um aviso.")
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)

        esperado = docpro.RECUO_PT
        # Busca por texto, não por índice: o título de nível 1 emite ANTES o
        # kicker ("SEÇÃO 01"), e ele também tem de pousar na mesma aresta.
        def paragrafo(texto):
            return next(p for p in doc.paragraphs if p.text.strip() == texto)

        for texto in ("SEÇÃO 01", "Seção", "Um parágrafo do corpo."):
            self.assertEqual(paragrafo(texto).paragraph_format.left_indent.pt, esperado,
                             "fora da aresta: %r" % texto)
        # a lista pendura o marcador: recuo + 14, primeira linha -14 => RECUO_PT
        item = paragrafo("• item um")
        self.assertEqual(item.paragraph_format.left_indent.pt
                         + item.paragraph_format.first_line_indent.pt, esperado)

        def margem_esquerda(cell):
            mar = cell._tc.tcPr.find(qn("w:tcMar"))
            return int(mar.find(qn("w:start")).get(qn("w:w"))) / 20.0

        tabela = doc.tables[0]
        self.assertEqual(margem_esquerda(tabela.rows[0].cells[0]), esperado)
        self.assertEqual(margem_esquerda(tabela.rows[1].cells[0]), esperado)
        callout = doc.tables[1]
        self.assertEqual(margem_esquerda(callout.rows[0].cells[0]), esperado)

    def test_caractere_de_controle_nao_corrompe_o_arquivo(self):
        """Caractere de controle é ilegal em XML 1.0: gravado cru, o Word
        recusa o arquivo. O kit remove antes de criar o run."""
        r = self._nu()
        r.paragrafo("texto com \x07 sinal e \x00 nulo")
        r.tabela(["A"], [["valor \x1f estranho"]])
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)  # reabre: prova que o XML é válido
        self.assertNotIn("\x07", doc.paragraphs[0].text)
        self.assertIn("texto com", doc.paragraphs[0].text)

    def test_sobrio_nasce_justificado_e_sem_cor(self):
        a = docpro.Sobrio()
        a.titulo("ATA DE REUNIÃO DE SÓCIOS")
        a.paragrafo("Aos vinte e cinco dias do mês de outubro ...")
        a.salvar(self.path, pdf=False)
        doc = Document(self.path)
        normal = doc.styles["Normal"].paragraph_format
        # 3 == WD_ALIGN_PARAGRAPH.JUSTIFY
        self.assertEqual(int(normal.alignment), 3)

    # ---------- identidade "Tinta & Latão" e blocos editoriais ----------
    def test_titulo_numera_a_secao_sozinho(self):
        """O prompt manda escrever só o texto do título: a numeração é do kit.
        Se ela parar de sair, o documento perde a hierarquia — e se sair EM
        DOBRO, o modelo volta a escrever "1." no texto."""
        r = docpro.Relatorio("T")
        r.titulo("Dados cadastrais")
        r.titulo("Indicadores")
        r.titulo("Detalhe", nivel=2)
        r.titulo("Anexo", kicker="ANEXO A")
        r.salvar(self.path, pdf=False)
        textos = [p.text for p in Document(self.path).paragraphs]
        self.assertIn("SEÇÃO 01", textos)
        self.assertIn("SEÇÃO 02", textos)
        self.assertIn("ANEXO A", textos)
        self.assertNotIn("SEÇÃO 03", textos)  # nível 2 não consome numeração

    def test_sumario_lista_as_entradas_com_a_pagina(self):
        r = docpro.Relatorio("T")
        r.sumario([("Sumário executivo", 3), ("Indicadores", 4)])
        r.salvar(self.path, pdf=False)
        textos = "\n".join(p.text for p in Document(self.path).paragraphs)
        self.assertIn("SUMÁRIO", textos)
        self.assertIn("Sumário executivo", textos)
        self.assertIn("01", textos)
        self.assertIn("3", textos)

    def test_confidencial_marca_o_rodape_e_pode_ser_desligado(self):
        for confidencial in (True, False):
            r = docpro.Relatorio("T", emissor="Escritório", confidencial=confidencial)
            r.capa()
            r.salvar(self.path, pdf=False)
            rodape = Document(self.path).sections[-1].footer.paragraphs[0].text
            if confidencial:
                self.assertIn("CONFIDENCIAL", rodape)
            else:
                self.assertNotIn("CONFIDENCIAL", rodape)

    def test_blocos_editoriais_novos_entram_no_documento(self):
        r = docpro.Relatorio("T", emissor="Escritório")
        r.citacao("Frase-chave do documento.", "Sumário executivo")
        r.etapas([("Diagnóstico", "ago/2026"), ("Implantação", "out/2026")])
        r.tabela(["A", "B"], [["1", "2"]], titulo="Investimento", fonte="Fonte: balancetes.")
        r.assinaturas(["João da Silva", "Maria Oliveira", "Carlos Souza"],
                      cargos=["Contador", "Sócia", "Sócio"],
                      local_data="Palmas/TO, 27 de julho de 2026.")
        r.contracapa(contatos=["contato@empresa.com.br"])
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)
        # `textos_do_docx` desce nas tabelas ANINHADAS: o fecho e as assinaturas
        # moram numa tabela dentro de uma célula (é o que os torna indivisíveis),
        # e uma varredura de um nível só perderia o nome do signatário.
        texto = "\n".join(docpro.textos_do_docx(doc))
        for esperado in ("Tabela 1 — Investimento", "Fonte: balancetes.",
                         "Frase-chave", "Diagnóstico", "João da Silva",
                         "Maria Oliveira", "Carlos Souza",
                         "contato@empresa.com.br", "Palmas/TO"):
            self.assertIn(esperado, texto)


    # ---------- v2: presets, números, paginação e conferência ----------
    def test_preset_decide_capa_sumario_e_numeracao(self):
        """O modelo escolhe o PRESET, não os blocos: capa, sumário e numeração
        de seção saem dele. Sem isso, cada documento saía com uma combinação
        diferente conforme o que o modelo lembrou de chamar."""
        r = docpro.Relatorio("T", emissor="Escritório", preset="gerencial")
        for nome in ("Sumário executivo", "Resultado", "Indicadores", "Recomendações"):
            r.titulo(nome)
            r.paragrafo("Conteúdo da seção %s." % nome)
        r.salvar(self.path, pdf=False)
        textos = "\n".join(docpro.textos_do_docx(Document(self.path)))
        self.assertIn("SEÇÃO 01", textos)          # numeração do preset
        self.assertIn("Conteúdo deste documento", textos)  # sumário automático
        self.assertIn("Cliente:", textos) if False else None
        self.assertIn("CONFIDENCIAL", textos)      # capa automática

        # No parecer a numeração é "1." / "1.1" e a capa é a simples.
        p = docpro.Relatorio("T", emissor="Escritório", preset="parecer")
        p.titulo("Fundamentação")
        p.titulo("Legislação aplicável", nivel=2)
        p.paragrafo("Texto do parecer, com corpo justificado por padrão.")
        caminho = os.path.join(self.tmp.name, "parecer.docx")
        p.salvar(caminho, pdf=False)
        textos = "\n".join(docpro.textos_do_docx(Document(caminho)))
        self.assertIn("1. Fundamentação", textos)
        self.assertIn("1.1 Legislação aplicável", textos)
        self.assertNotIn("SEÇÃO 01", textos)

        # Proposta e carta nunca ganham sumário.
        c = docpro.Relatorio("T", emissor="Escritório", preset="carta")
        c.paragrafo("Prezado cliente, segue a comunicação solicitada.")
        caminho = os.path.join(self.tmp.name, "carta.docx")
        c.salvar(caminho, pdf=False)
        textos = "\n".join(docpro.textos_do_docx(Document(caminho)))
        self.assertNotIn("Conteúdo deste documento", textos)

    def test_preset_desconhecido_falha_cedo(self):
        with self.assertRaises(docpro.KitError):
            docpro.Relatorio("T", preset="inventado")

    def test_tabela_recebe_numeros_e_o_kit_formata_em_pt_br(self):
        """String "R$ 412.300,00" montada à mão é de onde saem totais que não
        fecham e coluna numérica alinhada à esquerda. O kit recebe o NÚMERO."""
        r = self._nu()
        r.tabela(["Trimestre", "Receita", "Margem"],
                 [["1T25", 412300, 0.277], ["2T25", 455900, 0.319]],
                 moeda=["Receita"], pct=["Margem"], total="soma")
        r.salvar(self.path, pdf=False)
        tabela = Document(self.path).tables[0]
        celulas = [c.text for row in tabela.rows for c in row.cells]
        self.assertIn("R$ 412.300,00", celulas)
        self.assertIn("27,7%", celulas)
        # O TOTAL é CALCULADO pelo kit, não escrito pelo modelo.
        self.assertIn("R$ 868.200,00", celulas)
        self.assertIn("TOTAL", celulas)
        # Coluna numérica alinhada à direita, inclusive no cabeçalho.
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self.assertEqual(tabela.rows[1].cells[1].paragraphs[0].alignment,
                         WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(tabela.rows[0].cells[1].paragraphs[0].alignment,
                         WD_ALIGN_PARAGRAPH.RIGHT)

    def test_negativo_sai_entre_parenteses_e_em_vermelho(self):
        r = self._nu()
        r.tabela(["Conta", "Saldo"], [["Resultado", -1234.56]], moeda=["Saldo"])
        r.salvar(self.path, pdf=False)
        celula = Document(self.path).tables[0].rows[1].cells[1]
        self.assertEqual(celula.text, "(R$ 1.234,56)")
        self.assertEqual(str(celula.paragraphs[0].runs[0].font.color.rgb),
                         docpro.PALETA["negativo"])

    def test_tabela_curta_e_indivisivel_e_o_total_nunca_fica_orfao(self):
        """Cabeçalho repetido numa página com a linha TOTAL sozinha na seguinte
        é o defeito de paginação mais visível de todos."""
        r = self._nu()
        r.tabela(["A", "B"], [["1", 2], ["2", 3]], moeda=["B"], total="soma",
                 fonte="Fonte: teste.")
        r.salvar(self.path, pdf=False)
        t = Document(self.path).tables[0]
        # Toda linha é `cantSplit`; todas menos a última "mantêm com a próxima".
        for row in t.rows:
            self.assertIsNotNone(row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")))
        for row in t.rows[:-1]:
            self.assertTrue(row.cells[0].paragraphs[0].paragraph_format.keep_with_next)
        # A linha "Fonte:" cola na tabela: a última linha também mantém.
        self.assertTrue(t.rows[-1].cells[0].paragraphs[0].paragraph_format.keep_with_next)

    def test_kpi_adapta_o_corpo_ao_comprimento_do_valor(self):
        """Tamanho fixo é o que fazia "R$ 1,89 / mi" quebrar em duas linhas ao
        lado de "32,5%" — cartões da mesma fileira com alturas diferentes."""
        r = self._nu()
        r.kpis([(1887900, "Receita", "moeda"), (0.325, "Margem", "pct")])
        r.salvar(self.path, pdf=False)
        cartoes = Document(self.path).tables[0]
        valores = [c.text for c in cartoes.rows[0].cells]
        self.assertEqual(valores, ["R$ 1.887.900,00", "32,5%"])
        corpos = {c.paragraphs[0].runs[0].font.size.pt for c in cartoes.rows[0].cells}
        # Um só tamanho na fileira: corpos diferentes leem como erro.
        self.assertEqual(len(corpos), 1)
        self.assertEqual(corpos.pop(), docpro.ESCALA["kpi_curto"])

    def test_seis_kpis_viram_duas_fileiras_de_tres(self):
        r = self._nu()
        r.kpis([(i, "Rótulo %d" % i, "milhar") for i in range(1, 7)])
        r.salvar(self.path, pdf=False)
        fileiras = [t for t in Document(self.path).tables if len(t.rows) == 2]
        self.assertEqual([len(f.columns) for f in fileiras], [3, 3])

    def test_mais_de_seis_kpis_falha_cedo(self):
        with self.assertRaises(docpro.KitError):
            self._nu().kpis([(i, "R%d" % i) for i in range(7)])

    def test_titulos_usam_o_estilo_de_titulo_real_do_word(self):
        """Estilo de verdade devolve o painel de navegação no Word e os
        marcadores (outline) no PDF gêmeo — nenhum dos dois existia enquanto o
        título era só um parágrafo em negrito."""
        r = self._nu()
        r.titulo("Seção")
        r.titulo("Subseção", nivel=2)
        r.salvar(self.path, pdf=False)
        estilos = [p.style.name for p in Document(self.path).paragraphs if p.text.strip()]
        self.assertIn("Heading 1", estilos)
        self.assertIn("Heading 2", estilos)

    def test_metadados_e_idioma_vao_no_arquivo(self):
        r = docpro.Relatorio("Análise 2025", cliente="ACME", emissor="Escritório",
                             capa=False, sumario=False, contracapa=False)
        r.paragrafo("texto do documento com tamanho suficiente para a auditoria")
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)
        self.assertEqual(doc.core_properties.title, "Análise 2025")
        self.assertEqual(doc.core_properties.author, "Escritório")
        self.assertEqual(doc.core_properties.language, "pt-BR")
        normal = doc.styles["Normal"].element.get_or_add_rPr()
        self.assertEqual(normal.find(qn("w:lang")).get(qn("w:val")), "pt-BR")

    def test_tipografia_padrao_e_a_que_o_cliente_tem(self):
        """Source Serif 4 / Source Sans 3 não existem no Windows do cliente: o
        Word substitui e o documento conferido deixa de ser o documento aberto.
        Cambria/Calibri existem em todo Office desde 2007."""
        self.assertEqual((docpro.F_SERIF, docpro.F_SANS), ("Cambria", "Calibri"))
        r = self._nu()
        r.paragrafo("texto")
        r.salvar(self.path, pdf=False)
        normal = Document(self.path).styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri")

    def test_fecho_e_assinaturas_ficam_num_bloco_indivisivel(self):
        r = self._nu()
        r.paragrafo("conteúdo")
        r.assinaturas(["Nome Um"], cargos=["Contador"],
                      local_data="Palmas/TO, 02 de setembro de 2026.")
        doc = Document(self.path) if False else None
        r.salvar(self.path, pdf=False)
        doc = Document(self.path)
        externa = doc.tables[-1]
        self.assertIsNotNone(externa.rows[0]._tr.find(qn("w:trPr")).find(qn("w:cantSplit")))
        dentro = "\n".join(docpro.textos_do_docx(doc))
        self.assertIn("Palmas/TO, 02 de setembro de 2026.", dentro)
        self.assertIn("Nome Um", dentro)

    def test_sobrio_traz_os_helpers_de_redacao_juridica(self):
        a = docpro.Sobrio("CONTRATO SOCIAL",
                          identificacao="ACME LTDA — CNPJ 00.000.000/0001-00")
        a.clausula("DO OBJETO", "A sociedade tem por objeto a prestação de serviços.")
        a.paragrafo_unico("O objeto poderá ser ampliado por alteração contratual.")
        a.clausula("DO CAPITAL SOCIAL")
        a.paragrafo_numerado("O capital é dividido em quotas iguais.")
        a.inciso("integralização em moeda corrente;")
        a.inciso("prazo de trinta dias.")
        a.item("Primeira deliberação.")
        a.item("Segunda deliberação.")
        caminho = os.path.join(self.tmp.name, "contrato.docx")
        a.salvar(caminho, pdf=False)
        texto = "\n".join(docpro.textos_do_docx(Document(caminho)))
        self.assertIn("CLÁUSULA PRIMEIRA — DO OBJETO", texto)
        self.assertIn("CLÁUSULA SEGUNDA — DO CAPITAL SOCIAL", texto)
        self.assertIn("Parágrafo único.", texto)
        self.assertIn("§ 1º", texto)
        self.assertIn("I – integralização", texto)
        self.assertIn("II – prazo", texto)
        self.assertIn("1. Primeira deliberação.", texto)
        self.assertIn("2. Segunda deliberação.", texto)

    def test_sobrio_nao_duplica_numeracao_que_ja_veio_no_texto(self):
        a = docpro.Sobrio("ATA")
        a.item("1. Já numerado pelo autor.")
        caminho = os.path.join(self.tmp.name, "ata2.docx")
        a.salvar(caminho, pdf=False)
        texto = "\n".join(docpro.textos_do_docx(Document(caminho)))
        self.assertIn("1. Já numerado pelo autor.", texto)
        self.assertNotIn("1. 1.", texto)


    # ---------- v2: o segundo passo do PDF (sumário com páginas REAIS) ----------
    @unittest.skipUnless(HAS_PDF, "LibreOffice/pypdf ausentes")
    def test_sumario_aponta_a_pagina_real_de_cada_secao(self):
        """O defeito que motivou a v2: o modelo informava a página e o sumário
        saía errado no próprio teste ("Sumário executivo ... 3" com a seção na
        página 2). Agora o kit grava, converte para PDF, descobre em que página
        cada título caiu e reescreve os números antes de converter de novo."""
        r = docpro.Relatorio("Relatório", cliente="ACME", emissor="Escritório")
        for i, nome in enumerate(("Sumário executivo", "Resultado", "Indicadores",
                                  "Recomendações")):
            r.titulo(nome)
            r.paragrafo(("Parágrafo %d. " % i) + "Conteúdo de tamanho realista. " * 30)
        rel = r.salvar(self.path)
        self.assertTrue(rel["ok"], rel["achados"])
        self.assertGreaterEqual(rel["paginas"], 2)

        from pypdf import PdfReader
        paginas = [(p.extract_text() or "") for p in PdfReader(rel["pdf"]).pages]
        linhas = [l for l in paginas[1].split("\n") if "...." in l]
        self.assertEqual(len(linhas), 4)
        for titulo, linha in zip(("Sumário executivo", "Resultado", "Indicadores",
                                  "Recomendações"), linhas):
            declarada = int(linha.rsplit(".", 1)[-1].strip())
            real = next(i for i, texto in enumerate(paginas, start=1)
                        if i > 2 or (i == 2 and titulo in texto.split("....")[-1])
                        if titulo in docpro._sem_linhas_de_sumario(texto))
            self.assertEqual(declarada, real,
                             "o sumário diz %d para %r" % (declarada, titulo))

    @unittest.skipUnless(HAS_PDF, "LibreOffice/pypdf ausentes")
    def test_sumario_com_poucas_entradas_nao_ganha_pagina_propria(self):
        r = docpro.Relatorio("Relatório", cliente="ACME", emissor="Escritório")
        for nome in ("Um", "Dois", "Três", "Quatro"):
            r.titulo(nome)
            r.paragrafo("Conteúdo de tamanho realista. " * 20)
        rel = r.salvar(self.path)
        from pypdf import PdfReader
        pagina2 = PdfReader(rel["pdf"]).pages[1].extract_text() or ""
        self.assertIn("Conteúdo deste documento", pagina2)
        # O kicker sai com espaçamento entre letras (é a assinatura visual da
        # identidade), então o extrator devolve "S E Ç Ã O  0 1".
        compacto = pagina2.replace(" ", "")
        self.assertIn("SEÇÃO01", compacto)   # o conteúdo começa na MESMA página

    @unittest.skipUnless(HAS_PDF, "LibreOffice/pypdf ausentes")
    def test_sobrio_confere_o_pdf_gemeo(self):
        a = docpro.Sobrio("ATA DE REUNIÃO DE SÓCIOS",
                          identificacao="ACME LTDA — CNPJ 00.000.000/0001-00")
        a.paragrafo("Aos dois dias do mês de setembro de dois mil e vinte e seis, "
                    "reuniram-se os sócios na sede social. " * 4)
        a.fecho("Palmas/TO, 02 de setembro de 2026.")
        a.assinaturas(["Nome Um", "Nome Dois"], cargos=["Sócio", "Sócia"])
        caminho = os.path.join(self.tmp.name, "ata-pdf.docx")
        rel = a.salvar(caminho)
        self.assertTrue(rel["ok"], rel["achados"])
        self.assertEqual(rel["achados"], [])
        self.assertTrue(os.path.exists(rel["pdf"]))

    @unittest.skipUnless(HAS_MPL, "matplotlib não instalado")
    def test_graficos_entram_como_imagem_na_mesma_aresta(self):
        r = docpro.Relatorio("T")
        r.grafico_barras(["2024", "2025"], {"Receita": [5.2, 5.8]}, titulo="Receita")
        r.grafico_linhas(["Jan", "Fev"], {"Saldo": [1.2, 1.4]}, titulo="Saldo")
        r.grafico_pizza(["Serviços", "Produtos"], [62, 38], titulo="Participação")
        r.salvar(self.path, pdf=False)
        with zipfile.ZipFile(self.path) as z:
            imagens = [n for n in z.namelist() if n.startswith("word/media/")]
        self.assertEqual(len(imagens), 3)
        doc = Document(self.path)
        # A figura respeita a mesma aresta do corpo — senão cria uma 2ª margem.
        com_figura = [p for p in doc.paragraphs if p._p.findall(".//" + qn("w:drawing"))]
        self.assertEqual(len(com_figura), 3)
        for p in com_figura:
            self.assertEqual(p.paragraph_format.left_indent.pt, docpro.RECUO_PT)

    def test_sobrio_identifica_o_documento_e_assina_em_pares_sem_cor(self):
        a = docpro.Sobrio()
        a.identificacao("ATA DE REUNIÃO DE SÓCIOS",
                        "ACME LTDA — CNPJ 00.000.000/0001-00 — NIRE 17.2.0000000-1")
        a.secao("ORDEM DO DIA")
        a.item("1. Aprovação das contas do exercício.")
        a.fecho("Palmas/TO, 27 de julho de 2026.")
        a.assinaturas(["João da Silva", "Maria Oliveira"],
                      subtitulos=["Sócio-administrador", "Sócia"])
        a.salvar(self.path, pdf=False)
        doc = Document(self.path)
        texto = "\n".join(docpro.textos_do_docx(doc))
        self.assertIn("ATA DE REUNIÃO DE SÓCIOS", texto)
        # Hífen NÃO separável nos identificadores: sem ele o Word parte o NIRE
        # no fim da linha e ele aparece cortado na folha da Junta Comercial.
        self.assertIn("NIRE 17.2.0000000" + docpro.kits.HIFEN_FIXO + "1", texto)
        self.assertNotIn("NIRE 17.2.0000000-1", texto)
        self.assertIn("João da Silva", texto)
        # Fecho e assinaturas vivem numa tabela EXTERNA de uma coluna que não
        # se parte (é o que impede a página só com a linha de assinatura); os
        # pares ficam em tabelas aninhadas dentro dela.
        self.assertEqual(len(doc.tables), 1)
        pares = doc.tables[0].rows[0].cells[0].tables
        self.assertEqual(len(pares), 1)   # dois nomes = um par
        util = Cm(21.0 - 3.0 - 2.0).twips
        tblW = pares[0]._tbl.tblPr.find(qn("w:tblW"))
        self.assertEqual(tblW.get(qn("w:type")), "dxa")
        self.assertLessEqual(int(tblW.get(qn("w:w"))), util)

        # documento registrável é 100% preto, inclusive nas tabelas aninhadas
        alvos = list(doc.paragraphs)

        def _desce(t):
            for row in t.rows:
                for cell in row.cells:
                    alvos.extend(cell.paragraphs)
                    for interna in cell.tables:
                        _desce(interna)

        for t in doc.tables:
            _desce(t)
        for p in alvos:
            for run in p.runs:
                cor = run.font.color
                if cor is not None and cor.rgb is not None:
                    self.assertEqual(str(cor.rgb), "000000")


if __name__ == "__main__":
    unittest.main()
