"""Testes do validar_artefato — F-23 da auditoria.

O que estes testes protegem é a frase "arquivo verificado". O validador decide
se a entrega se apresenta como boa; até agora ele era Python dentro de uma
template string de JavaScript, e **nenhum teste o alcançava**. Um validador sem
teste é pior que validador nenhum: sem ele a entrega diz "não verifiquei", com
ele quebrado a entrega diz "verificado" — e essa é a mentira que o F-10 já
tinha custado caro.

Todos os arquivos aqui são REAIS: .xlsx escrito com openpyxl, .docx com
python-docx, .pdf com reportlab. Nada de dublê — o que se quer provar é
justamente que o validador lida com o que as bibliotecas produzem de verdade.

Os gráficos quebrados são a parte que exige cirurgia: openpyxl não deixa
escrever uma referência inválida, então geramos o gráfico bom e reescrevemos o
`xl/charts/chart1.xml` dentro do zip. É como o defeito aparece no mundo real —
o modelo gera a planilha e a referência não bate com as abas.
"""
import os
import re
import shutil
import tempfile
import unittest
import zipfile

try:
    from docx import Document
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, Reference
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdfcanvas

    import validar_artefato as va

    TEM_LIBS = True
except Exception:  # libs ausentes no ambiente — o CI e o sandbox as instalam
    TEM_LIBS = False


def sem_recalculo(**kw):
    """Config de teste: recálculo DESLIGADO.

    O recálculo chama o `soffice`, que não existe (nem deve existir) no runner
    dos testes. Deixá-lo ligado tornaria a suíte dependente do LibreOffice e
    lenta por 25 s de timeout a cada planilha. O caminho do recálculo é
    exercitado separadamente em `test_recalculo_*`, com o subprocess encurtado.
    """
    kw.setdefault("recalc_enabled", False)
    return va.Config(**kw)


@unittest.skipUnless(TEM_LIBS, "requer python-docx, openpyxl e reportlab")
class BaseArtefato(unittest.TestCase):
    def setUp(self):
        self.dir = tempfile.mkdtemp(prefix="valida_")
        self.addCleanup(shutil.rmtree, self.dir, ignore_errors=True)

    def caminho(self, nome):
        return os.path.join(self.dir, nome)

    def planilha(self, nome="p.xlsx", celulas=None, abas=("Dados",)):
        wb = Workbook()
        wb.remove(wb.active)
        for aba in abas:
            wb.create_sheet(aba)
        for (aba, coord, valor) in (celulas or []):
            wb[aba][coord] = valor
        p = self.caminho(nome)
        wb.save(p)
        return p

    def planilha_com_grafico(self, nome="g.xlsx", valores=(1, 2, 3), ref_xml=None):
        """Planilha com um gráfico de barras de verdade.

        `ref_xml` substitui a referência da série depois de salvar — é como
        injetamos aba inexistente, intervalo invertido e série vazia.
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Dados"
        ws["A1"] = "Mes"
        ws["B1"] = "Total"
        for i, v in enumerate(valores, start=2):
            ws["A" + str(i)] = "M" + str(i - 1)
            if v is not None:
                ws["B" + str(i)] = v
        ch = BarChart()
        ch.add_data(Reference(ws, min_col=2, min_row=1, max_row=1 + len(valores)), titles_from_data=True)
        ch.set_categories(Reference(ws, min_col=1, min_row=2, max_row=1 + len(valores)))
        ws.add_chart(ch, "D2")
        p = self.caminho(nome)
        wb.save(p)
        if ref_xml:
            self.reescrever_ref_do_grafico(p, ref_xml)
        return p

    def reescrever_ref_do_grafico(self, p, nova_ref):
        """Troca a referência da SÉRIE DE VALORES dentro do zip do .xlsx.

        Mexe só no bloco `<val>` — o `<cat>` (categorias) fica intacto, porque
        os defeitos que queremos injetar são de série de valores. O prefixo de
        namespace é opcional na regex: o openpyxl escreve `<val>` sem prefixo,
        mas o Excel escreve `<c:val>`, e o validador aceita os dois.
        """
        alvo = "xl/charts/chart1.xml"
        with zipfile.ZipFile(p) as z:
            itens = {n: z.read(n) for n in z.namelist()}
        xml = itens[alvo].decode("utf-8")

        def troca(m):
            return re.sub(r"(<(?:\w+:)?f>)[^<]+(</(?:\w+:)?f>)", r"\g<1>" + nova_ref + r"\g<2>", m.group(0))

        novo = re.sub(r"<(?:\w+:)?val>.*?</(?:\w+:)?val>", troca, xml, flags=re.S)
        self.assertNotEqual(novo, xml, "o helper precisa achar o bloco <val> — se falhar aqui, o teste não prova nada")
        itens[alvo] = novo.encode("utf-8")
        with zipfile.ZipFile(p, "w", zipfile.ZIP_DEFLATED) as z:
            for n, b in itens.items():
                z.writestr(n, b)

    def documento(self, nome="d.docx", paragrafos=(), com_tabela=False):
        d = Document()
        for t in paragrafos:
            d.add_paragraph(t)
        if com_tabela:
            d.add_table(rows=2, cols=2)
        p = self.caminho(nome)
        d.save(p)
        return p

    def pdf(self, nome="a.pdf", paginas=1):
        p = self.caminho(nome)
        c = pdfcanvas.Canvas(p, pagesize=A4)
        for i in range(paginas):
            c.drawString(100, 700, "pagina " + str(i + 1))
            c.showPage()
        c.save()
        return p


class TestPlanilha(BaseArtefato):
    def test_planilha_limpa_passa_e_conta_as_abas(self):
        p = self.planilha(abas=("Dados", "Resumo"), celulas=[("Dados", "A1", 10)])
        r = va.check_xlsx(p, sem_recalculo())
        self.assertTrue(r["ok"])
        self.assertIn("2 abas", r["info"])

    def test_celula_com_erro_de_formula_reprova(self):
        # O caso que motiva o validador: a planilha ABRE, tem a cara certa e
        # traz #REF! numa célula. Sem esta checagem a entrega diria "ok".
        p = self.planilha(celulas=[("Dados", "A1", "Total"), ("Dados", "B1", "#REF!")])
        r = va.check_xlsx(p, sem_recalculo())
        self.assertFalse(r["ok"])
        self.assertIn("1 celula(s) com erro de formula", r["info"])

    def test_reconhece_todos_os_codigos_de_erro_do_excel(self):
        # Cada código é um modo de falha distinto; um só na lista já bastaria
        # para o teste passar sem cobrir os outros.
        for codigo in va.ERRS:
            with self.subTest(codigo=codigo):
                p = self.planilha(nome="e.xlsx", celulas=[("Dados", "A1", codigo)])
                self.assertFalse(va.check_xlsx(p, sem_recalculo())["ok"])

    def test_erro_dentro_de_texto_maior_tambem_conta(self):
        p = self.planilha(celulas=[("Dados", "A1", "Resultado: #DIV/0! (conferir)")])
        self.assertFalse(va.check_xlsx(p, sem_recalculo())["ok"])

    def test_sem_recalculo_a_verificacao_se_declara_parcial(self):
        # Honestidade do relatório: sem recalcular, o validador NÃO pode dizer
        # que as fórmulas estão boas — só que não achou erro no texto delas.
        p = self.planilha(celulas=[("Dados", "A1", "=1/0")])
        r = va.check_xlsx(p, sem_recalculo())
        self.assertIn("formulas NAO recalculadas (verificacao parcial)", r["info"])

    def test_teto_de_celulas_e_declarado_no_relatorio(self):
        # Corte silencioso conta história falsa com número verdadeiro: se a
        # varredura parou na 5ª célula, o relatório tem de dizer isso.
        celulas = [("Dados", "A" + str(i), i) for i in range(1, 30)]
        p = self.planilha(celulas=celulas)
        r = va.check_xlsx(p, sem_recalculo(max_cells=5))
        self.assertIn("varredura limitada a 5 celulas", r["info"])

    def test_erro_alem_do_teto_passa_despercebido_e_isso_e_declarado(self):
        # Consequência assumida do teto — o teste existe para que ninguém leia
        # "ok" como "sem erro na planilha inteira".
        celulas = [("Dados", "A" + str(i), i) for i in range(1, 20)]
        celulas.append(("Dados", "A50", "#REF!"))
        p = self.planilha(celulas=celulas)
        r = va.check_xlsx(p, sem_recalculo(max_cells=3))
        self.assertTrue(r["ok"])
        self.assertIn("varredura limitada", r["info"])

    def test_xlsm_segue_o_mesmo_caminho_do_xlsx(self):
        p = self.planilha(nome="m.xlsm", celulas=[("Dados", "A1", "#N/A")])
        r = va.validar(["m.xlsm"], base=self.dir, cfg=sem_recalculo())
        self.assertFalse(r[0]["ok"])


class TestGraficos(BaseArtefato):
    def test_grafico_saudavel_e_contado(self):
        p = self.planilha_com_grafico()
        r = va.check_xlsx(p, sem_recalculo())
        self.assertTrue(r["ok"], r["info"])
        self.assertIn("1 grafico(s) ok", r["info"])

    def test_grafico_apontando_para_aba_inexistente_reprova(self):
        p = self.planilha_com_grafico(ref_xml="'NaoExiste'!$B$1:$B$4")
        r = va.check_xlsx(p, sem_recalculo())
        self.assertFalse(r["ok"])
        self.assertIn("aba inexistente", r["info"])

    def test_grafico_com_intervalo_invertido_reprova(self):
        # C2:B2 abre vazio no Excel. É o defeito que passava despercebido
        # porque o arquivo ABRE — só o gráfico é que fica em branco.
        p = self.planilha_com_grafico(ref_xml="'Dados'!$C$2:$B$2")
        r = va.check_xlsx(p, sem_recalculo())
        self.assertFalse(r["ok"])
        self.assertIn("intervalo invertido", r["info"])

    def test_grafico_com_serie_de_valores_vazia_reprova(self):
        # Coluna declarada no cabeçalho e deixada sem dados pelo modelo. Aqui
        # não há cirurgia no XML: o openpyxl gera a referência normalmente e as
        # células é que estão vazias — exatamente como o defeito nasce.
        p = self.planilha_com_grafico(valores=(None, None, None))
        r = va.check_xlsx(p, sem_recalculo())
        self.assertFalse(r["ok"])
        self.assertIn("serie de valores vazia", r["info"])

    def test_categoria_de_texto_nao_e_acusada_de_serie_vazia(self):
        # A contraprova do teste acima: a coluna de CATEGORIAS é texto por
        # natureza. Se o validador cobrasse número dela, todo gráfico normal
        # seria reprovado — o falso positivo que inutiliza a checagem.
        p = self.planilha_com_grafico(valores=(5, 6, 7))
        r = va.check_xlsx(p, sem_recalculo())
        self.assertTrue(r["ok"], r["info"])

    def test_planilha_sem_grafico_nao_inventa_problema(self):
        p = self.planilha(celulas=[("Dados", "A1", 1)])
        r = va.check_xlsx(p, sem_recalculo())
        self.assertTrue(r["ok"])
        self.assertNotIn("grafico", r["info"])


class TestReferencias(BaseArtefato):
    """A gramática das referências, direto nas funções puras."""

    def test_parse_de_referencia_com_e_sem_aspas(self):
        self.assertEqual(va._parse_ref("'Meus Dados'!$A$1:$B$2"), ("Meus Dados", (1, 1, 2, 2)))
        self.assertEqual(va._parse_ref("Dados!A1:B2"), ("Dados", (1, 1, 2, 2)))

    def test_referencia_de_celula_unica_vira_intervalo_degenerado_valido(self):
        self.assertEqual(va._parse_ref("Dados!$C$3"), ("Dados", (3, 3, 3, 3)))

    def test_referencia_sem_aba_nao_e_parseada(self):
        self.assertIsNone(va._parse_ref("A1:B2"))

    def test_colunas_de_duas_letras(self):
        self.assertEqual(va._col_to_num("AA"), 27)
        self.assertEqual(va._col_to_num("B"), 2)

    def test_intervalo_com_formula_conta_como_tendo_numero(self):
        # Fórmula ainda não calculada é promessa de número — acusá-la de vazia
        # reprovaria toda planilha gerada com fórmulas, que é a maioria.
        wb = Workbook()
        ws = wb.active
        ws.title = "Dados"
        ws["B2"] = "=SUM(A1:A3)"
        self.assertTrue(va._range_has_number(wb, "Dados", (2, 2, 2, 2)))

    def test_aba_inexistente_nao_acusa_por_desencargo(self):
        wb = Workbook()
        self.assertTrue(va._range_has_number(wb, "Fantasma", (1, 1, 2, 2)))


class TestDocumento(BaseArtefato):
    def test_documento_com_texto_passa(self):
        p = self.documento(paragrafos=("Olá", "mundo"))
        r = va.check_docx(p)
        self.assertTrue(r["ok"])
        self.assertIn("2 paragrafos", r["info"])

    def test_documento_vazio_reprova_e_diz_por_que(self):
        p = self.documento()
        r = va.check_docx(p)
        self.assertFalse(r["ok"])
        self.assertIn("documento vazio", r["info"])

    def test_documento_so_com_tabela_nao_e_vazio(self):
        # Relatório que é uma tabela e nada mais é entrega legítima. Contar só
        # parágrafos reprovaria um documento perfeitamente bom.
        p = self.documento(com_tabela=True)
        r = va.check_docx(p)
        self.assertTrue(r["ok"])
        self.assertIn("1 tabelas", r["info"])

    def test_paragrafos_so_com_espaco_nao_contam_como_conteudo(self):
        p = self.documento(paragrafos=("", "   ", "\t"))
        self.assertFalse(va.check_docx(p)["ok"])


class TestPdf(BaseArtefato):
    def test_pdf_de_uma_pagina(self):
        r = va.check_pdf(self.pdf())
        self.assertTrue(r["ok"])
        self.assertIn("1 paginas", r["info"])

    def test_pdf_de_varias_paginas_conta_certo(self):
        r = va.check_pdf(self.pdf(nome="v.pdf", paginas=4))
        self.assertIn("4 paginas", r["info"])


class TestRoteamento(BaseArtefato):
    def test_roteia_cada_extensao_para_o_validador_certo(self):
        self.planilha(nome="p.xlsx", celulas=[("Dados", "A1", 1)])
        self.documento(nome="d.docx", paragrafos=("texto",))
        self.pdf(nome="a.pdf")
        r = {x["path"]: x for x in va.validar(["p.xlsx", "d.docx", "a.pdf"], base=self.dir, cfg=sem_recalculo())}
        self.assertIn("abas", r["p.xlsx"]["info"])
        self.assertIn("paragrafos", r["d.docx"]["info"])
        self.assertIn("paginas", r["a.pdf"]["info"])
        self.assertTrue(all(x["ok"] for x in r.values()))

    def test_extensao_desconhecida_passa_sem_alegar_nada(self):
        # "Não sei julgar" tem de sair como ok+info vazio, nunca como reprovado:
        # reprovar o que não se sabe ler é inventar defeito.
        with open(self.caminho("notas.txt"), "w") as f:
            f.write("oi")
        r = va.validar(["notas.txt"], base=self.dir, cfg=sem_recalculo())
        self.assertTrue(r[0]["ok"])
        self.assertEqual(r[0]["info"], "")

    def test_arquivo_corrompido_reprova_dizendo_que_nao_abre(self):
        with open(self.caminho("quebrado.xlsx"), "wb") as f:
            f.write(b"isto nao e um zip")
        r = va.validar(["quebrado.xlsx"], base=self.dir, cfg=sem_recalculo())
        self.assertFalse(r[0]["ok"])
        self.assertIn("nao abre", r[0]["info"])

    def test_arquivo_inexistente_reprova_sem_derrubar_o_lote(self):
        # Um arquivo ruim não pode levar os outros junto: o laço tem de seguir.
        self.documento(nome="bom.docx", paragrafos=("texto",))
        r = va.validar(["sumiu.xlsx", "bom.docx"], base=self.dir, cfg=sem_recalculo())
        self.assertFalse(r[0]["ok"])
        self.assertTrue(r[1]["ok"])

    def test_mensagem_de_erro_e_truncada(self):
        with open(self.caminho("x" * 200 + ".docx"), "wb") as f:
            f.write(b"nao e docx")
        r = va.validar(["x" * 200 + ".docx"], base=self.dir, cfg=sem_recalculo())
        self.assertLessEqual(len(r[0]["info"]), 120)

    def test_a_ordem_da_resposta_segue_a_ordem_do_pedido(self):
        # Quem consome casa por `path`, mas a ordem estável é o que torna o
        # relatório legível quando dois arquivos têm o mesmo veredito.
        self.documento(nome="a.docx", paragrafos=("t",))
        self.documento(nome="b.docx", paragrafos=("t",))
        r = va.validar(["b.docx", "a.docx"], base=self.dir, cfg=sem_recalculo())
        self.assertEqual([x["path"] for x in r], ["b.docx", "a.docx"])

    def test_base_diferente_de_workspace_e_respeitada(self):
        # O default é /workspace porque é lá que o sandbox monta os arquivos;
        # o parâmetro existe para o teste — e é o que torna tudo isto possível.
        self.documento(nome="d.docx", paragrafos=("t",))
        self.assertTrue(va.validar(["d.docx"], base=self.dir, cfg=sem_recalculo())[0]["ok"])
        self.assertFalse(va.validar(["d.docx"], base="/nao/existe", cfg=sem_recalculo())[0]["ok"])


class TestConfig(unittest.TestCase):
    """A configuração lida do ambiente — os limites que o operador ajusta."""

    def test_valores_padrao(self):
        c = va.Config.from_env(env={})
        self.assertEqual(c.max_cells, 40000)
        self.assertTrue(c.recalc_enabled)
        self.assertEqual(c.recalc_timeout, 25)
        self.assertEqual(c.budget, [2])

    def test_ambiente_sobrepoe_os_padroes(self):
        c = va.Config.from_env(env={
            "VALIDATE_MAX_CELLS": "10",
            "VALIDATE_RECALC": "false",
            "SANDBOX_RECALC_TIMEOUT_S": "3",
            "RECALC_MAX_FILES": "1",
        })
        self.assertEqual(c.max_cells, 10)
        self.assertFalse(c.recalc_enabled)
        self.assertEqual(c.recalc_timeout, 3)
        self.assertEqual(c.budget, [1])

    def test_recalc_desligado_aceita_maiusculas(self):
        self.assertFalse(va.Config.from_env(env={"VALIDATE_RECALC": "FALSE"}).recalc_enabled)

    def test_orcamento_e_por_chamada_nao_global(self):
        # Era um global mutável: uma execução gastava o orçamento da seguinte.
        # Duas Config novas têm de nascer com o mesmo saldo.
        self.assertEqual(va.Config().budget, va.Config().budget)


@unittest.skipUnless(TEM_LIBS, "requer python-docx, openpyxl e reportlab")
class TestOrcamentoDeRecalculo(BaseArtefato):
    def test_recalculo_gasta_o_orcamento_e_para(self):
        # Sem `soffice` no runner, `recalc` devolve (None, tmp) e o validador
        # segue como "não recalculado" — que é justamente o comportamento que
        # queremos provar: falta de LibreOffice não derruba a validação.
        cfg = va.Config(recalc_enabled=True, recalc_timeout=1, recalc_max_files=1)
        p = self.planilha(celulas=[("Dados", "A1", 1)])
        r1 = va.check_xlsx(p, cfg)
        self.assertTrue(r1["ok"])
        self.assertEqual(cfg.budget, [0], "a primeira planilha consome o orçamento")
        r2 = va.check_xlsx(p, cfg)
        self.assertIn("NAO recalculadas", r2["info"])


if __name__ == "__main__":
    unittest.main()
