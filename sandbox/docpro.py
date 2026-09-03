"""docpro — kit de DESIGN para documentos Word (.docx) com python-docx.

Instalado no sandbox do Frederico AI Studio, na identidade **"Tinta & Latão"**
(verde-tinta #0C3A30 com acento em latão #A9812F). A paleta, a escala
tipográfica, a formatação pt-BR e a auditoria vêm de `kits.py` — os três kits
partilham a mesma base, então Word, Excel e PDF do mesmo pacote formatam igual.

Três decisões estruturais que valem para todo bloco:

  * **a fonte que o cliente vê é a que foi conferida.** O padrão é Cambria
    (títulos e números de destaque) sobre Calibri (corpo e tabelas): existem em
    todo Office desde 2007 e, no Linux, o LibreOffice as substitui por
    Caladea/Carlito, **metricamente idênticas**. Assim o PDF gêmeo gerado aqui
    quebra a linha no mesmo lugar que o Word do cliente — a conferência passa a
    valer para o documento que ele realmente abre;
  * **uma aresta esquerda só** (`RECUO_PT`) — parágrafo, título, primeira
    coluna da tabela e conteúdo do callout pousam na mesma vertical;
  * **nenhuma tabela sem largura declarada** — 100% da área útil (`_fit`) ou
    largura fixa em dxa com a grade coerente (`_fixa`). É o que impede o
    clássico "tabela vazando a margem".

O modelo escolhe o PRESET, não os blocos: capa, sumário, numeração de seção,
alinhamento do corpo e fechamento saem dele (ver `PRESETS`).

Uso mínimo:
    from docpro import Relatorio
    r = Relatorio("Análise Econômico-Financeira 2025", cliente="ACME LTDA",
                  emissor="Meu Escritório", preset="gerencial")
    r.titulo("Sumário executivo")               # numera "SEÇÃO 01" sozinho
    r.paragrafo("Texto de análise ...")
    r.tabela(["Trimestre", "Receita"], [["1T25", 412300]],
             moeda=["Receita"], total="soma")   # NÚMEROS, não strings
    r.kpis([(1887900, "Receita líquida", "moeda")])
    rel = r.salvar("/workspace/outputs/relatorio.docx")
    print("CONFERÊNCIA:", rel)                  # {"ok": True, "paginas": 4, ...}

Documento registrável (ata, contrato, alteração contratual — JUCETINS) usa a
outra classe, `Sobrio`: ZERO cor, Times 12, corpo justificado de fábrica e
helpers de redação jurídica (`clausula`, `inciso`, `paragrafo_unico`).
"""
import os
import re
import subprocess
import tempfile
import warnings
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import kits
from kits import (ESCALA, KitError, PALETA, achado, achados_de_placeholder,
                  e_numero, escala_kpi, falha_se_grave, fmt, formata_valor,
                  limpa_texto, linha_de_total, linhas_de_kpi,
                  liga_identificadores, normaliza_linhas, paleta_para,
                  relatorio, tipos_de_coluna)

#: Tipografia com fidelidade no cliente (ver docstring do módulo).
F_SERIF = kits.TIPOGRAFIA["office"]["serif"]
F_SANS = kits.TIPOGRAFIA["office"]["sans"]
#: Cores dos gráficos, na ordem das séries/fatias.
CORES_GRAF = kits.CORES_GRAF

# ---------------------------------------------------------------------------
# Grade: existe uma única distância entre a margem do papel e o começo do texto,
# e todo bloco a respeita — parágrafo, título com barra, primeira coluna da
# tabela e conteúdo do callout. Sem isso o título nasce 10 pt à direita do corpo
# e a tabela 4 pt à esquerda dele: três arestas na mesma página, que é
# exatamente o que faz o documento parecer mal feito.
# ---------------------------------------------------------------------------

#: Recuo do texto em pontos. Em Word a margem interna de célula é medida em
#: dxa (1/20 de ponto), daí a conversão.
RECUO_PT = 10
RECUO_DXA = RECUO_PT * 20
#: Respiro interno das demais colunas da tabela (não afeta a aresta esquerda).
RESPIRO_DXA = 120

#: Acima deste número de linhas a tabela pode quebrar entre páginas (abaixo ela
#: é indivisível). Ver `_paginacao_tabela`.
LIMITE_TABELA_INDIVISIVEL = 15
#: Sumário com até este número de entradas nunca ganha página própria.
LIMITE_SUMARIO_SEM_PAGINA = 10

#: Registro por tipo de documento: o modelo escolhe o preset e o kit decide
#: capa, sumário, numeração, alinhamento do corpo e fechamento.
PRESETS = {
    "gerencial": {"capa": "faixa", "sumario": "auto", "numeracao": "secao",
                  "corpo": "esquerda", "fechamento": "auto",
                  "tipo": "RELATÓRIO GERENCIAL"},
    "parecer": {"capa": "simples", "sumario": "auto", "numeracao": "decimal",
                "corpo": "justificado", "fechamento": "faixa",
                "tipo": "PARECER TÉCNICO"},
    "proposta": {"capa": "faixa", "sumario": "nunca", "numeracao": None,
                 "corpo": "esquerda", "fechamento": "faixa",
                 "tipo": "PROPOSTA COMERCIAL"},
    "carta": {"capa": None, "sumario": "nunca", "numeracao": None,
              "corpo": "justificado", "fechamento": "assinatura",
              "tipo": "CARTA"},
}
#: Sumário automático só a partir deste número de seções de 1º nível.
MINIMO_SECOES_PARA_SUMARIO = 4


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


# Ordem canônica dos filhos de <w:tblPr> (CT_TblPr). O Word até tolera fora de
# ordem, mas python-docx e validadores estritos não — inserir tblW/tblBorders/
# tblLayout na posição certa evita um .docx que "abre corrompido" ou perde o
# estilo. Só listamos os elementos que este kit realmente usa.
_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
]


def _tblpr_put(tblPr, el):
    """Substitui/insere `el` em `tblPr` respeitando a ordem do schema OOXML."""
    tag = el.tag.split("}")[-1]
    for existente in tblPr.findall(qn("w:" + tag)):
        tblPr.remove(existente)
    posteriores = set(_TBLPR_ORDER[_TBLPR_ORDER.index(tag) + 1:]) if tag in _TBLPR_ORDER else set()
    ref = None
    for filho in tblPr:
        if filho.tag.split("}")[-1] in posteriores:
            ref = filho
            break
    if ref is None:
        tblPr.append(el)
    else:
        ref.addprevious(el)


def _shade(elem, fill):
    elem.append(_el("w:shd", val="clear", fill=fill))


def sombrear_celula(cell, cor):
    _shade(cell._tc.get_or_add_tcPr(), cor)


def _limpa_paragrafo(p):
    """Esvazia o parágrafo SEM deixar um run vazio para trás.

    `p.text = ""` (o caminho óbvio) apaga o conteúdo e acrescenta um run com
    texto vazio: `runs[0]` passa a ser esse fantasma, e quem for ler a fonte ou
    a cor da célula — a auditoria do documento sóbrio, por exemplo — lê um run
    sem formatação nenhuma em vez do texto de verdade."""
    for run in list(p.runs):
        p._p.remove(run._r)
    return p


def _cell_pad(cell, t=80, b=80, l=RESPIRO_DXA, r=RESPIRO_DXA):
    tcPr = cell._tc.get_or_add_tcPr()
    m = _el("w:tcMar")
    for k, v in (("top", t), ("bottom", b), ("start", l), ("end", r)):
        m.append(_el("w:" + k, w=v, type="dxa"))
    tcPr.append(m)


def _altura_linha(row, cm_, exata=True):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(_el("w:trHeight", val=int(Cm(cm_).pt * 20),
                    hRule="exact" if exata else "atLeast"))


def _nao_quebrar_linha(row):
    """`w:cantSplit`: a LINHA não se parte entre duas páginas."""
    row._tr.get_or_add_trPr().append(_el("w:cantSplit"))


def _colar_na_proxima(row, colar=True):
    """Em Word quem tem "manter com o próximo" é o PARÁGRAFO, não a linha da
    tabela — marcar os parágrafos de uma linha é o que faz ela viajar junto com
    a linha seguinte. É assim que a linha TOTAL nunca fica sozinha no topo de
    uma página."""
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = colar


def _num_texto(v):
    """O texto parece número (para alinhar a coluna à direita)?"""
    s = str(v).strip().replace("R$", "").replace(".", "").replace(",", ".").replace("%", "")
    s = s.replace("(", "-").replace(")", "").strip()
    if not s or s in "-+":
        return False
    try:
        float(s)
        return True
    except Exception:
        return False


def _run(p, texto, bold=False, size=ESCALA["corpo"], cor=None, branco=False,
         caps=False, spacing=None, serif=False, italico=False, fonte=None):
    r = p.add_run(limpa_texto(texto))
    escolhida = fonte or (F_SERIF if serif else F_SANS)
    r.font.name = escolhida
    rpr = r._element.get_or_add_rPr()
    # `font.name` sozinho só define o script latino: sem o w:rFonts completo, o
    # Word troca a fonte em qualquer caractere fora do ASCII.
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _el("w:rFonts")
        rpr.append(rfonts)
    for a in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn("w:" + a), escolhida)
    r.bold = bold
    r.italic = italico
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string("FFFFFF" if branco else (cor or PALETA["corpo"]))
    if caps:
        r.font.all_caps = True
    if spacing:
        rpr.append(_el("w:spacing", val=spacing))
    return r


def _idioma(estilo, lang="pt-BR"):
    """Declara o idioma no estilo — sem isto o Word revisa o documento inteiro
    com o corretor da língua da instalação e sublinha tudo de vermelho."""
    rpr = estilo.element.get_or_add_rPr()
    lang_el = rpr.find(qn("w:lang"))
    if lang_el is None:
        lang_el = _el("w:lang")
        rpr.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)


def _fonte_do_estilo(estilo, nome):
    estilo.font.name = nome
    rpr = estilo.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _el("w:rFonts")
        rpr.append(rfonts)
    for a in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn("w:" + a), nome)


def textos_do_docx(doc):
    """TODO o texto do documento: parágrafos, células e **tabelas aninhadas**.

    A recursão não é preciosismo: o bloco de fecho/assinaturas mora numa tabela
    dentro de uma célula, e uma varredura de um nível só deixaria justamente o
    nome do signatário e o CPF fora da auditoria de placeholder."""
    saida = [p.text for p in doc.paragraphs]

    def _tabela(t):
        for row in t.rows:
            for cell in row.cells:
                saida.extend(p.text for p in cell.paragraphs)
                for interna in cell.tables:
                    _tabela(interna)

    for t in doc.tables:
        _tabela(t)
    return saida


def _elementos_do_corpo(doc):
    """Filhos do corpo, sem o `w:sectPr` final (que não é conteúdo)."""
    return [e for e in doc.element.body if not e.tag.endswith("}sectPr")]


def _converter_pdf(caminho, timeout=120):
    """Converte para PDF com o LibreOffice. Devolve o caminho do PDF ou None.

    Silêncio aqui seria caro: é o PDF que a auditoria lê para achar sumário
    divergente, assinatura órfã e KPI quebrado. Quando a conversão não acontece,
    quem chama registra um AVISO em vez de fingir que conferiu."""
    saida = os.path.splitext(caminho)[0] + ".pdf"
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir",
             os.path.dirname(caminho) or ".", caminho],
            timeout=timeout, capture_output=True, check=False)
    except Exception:
        return None
    return saida if os.path.exists(saida) else None


def _texto_por_pagina(caminho_pdf):
    """Texto de cada página do PDF, na ordem. `[]` quando não dá para ler."""
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            return []
    try:
        return [(p.extract_text() or "") for p in PdfReader(caminho_pdf).pages]
    except Exception:
        return []


#: Linha de sumário: título, guia pontilhada e o número da página. O número é
#: OPCIONAL de propósito — no primeiro passo ele ainda não existe, e é
#: justamente aí que a busca precisa ignorar essas linhas.
_LINHA_DE_SUMARIO = re.compile(r"^.*\.{4,}\s*\d*\s*$", re.MULTILINE)


def _sem_linhas_de_sumario(texto):
    """Tira do texto da página as LINHAS do próprio sumário.

    Sem isso a busca pela página de uma seção acerta o sumário — que lista
    todos os títulos — e o índice inteiro passa a apontar para si mesmo. Foi
    esse o defeito do relatório que motivou a v2."""
    return _LINHA_DE_SUMARIO.sub("", texto)


def _pagina_do_texto(paginas, alvo, a_partir_de=0, ignorar_sumario=False):
    """Primeira página (1-based) cujo texto contém `alvo`. 0 = não achou."""
    chave = re.sub(r"\s+", " ", limpa_texto(alvo)).strip().lower()
    if not chave:
        return 0
    for i in range(max(0, a_partir_de), len(paginas)):
        texto = _sem_linhas_de_sumario(paginas[i]) if ignorar_sumario else paginas[i]
        if chave in re.sub(r"\s+", " ", texto).lower():
            return i + 1
    return 0


class Relatorio:
    """Documento Word montado só por blocos do kit.

    `preset=` decide capa, sumário, numeração de seção, alinhamento do corpo e
    fechamento (ver `PRESETS`). Qualquer um pode ser sobrescrito
    explicitamente: `capa=False`, `sumario=False`, `contracapa="faixa"`.
    """

    def __init__(self, titulo, cliente="", emissor="", subtitulo="",
                 tipo=None, data_str=None, margem_cm=2.0, cor_marca=None,
                 confidencial=True, preset="gerencial", tipografia="office",
                 capa=None, sumario=None, contracapa=None, idioma="pt-BR"):
        if preset not in PRESETS:
            raise KitError("preset desconhecido: %r (use %s)"
                           % (preset, " | ".join(sorted(PRESETS))))
        self.preset = preset
        self.regras = dict(PRESETS[preset])
        if tipografia not in kits.TIPOGRAFIA:
            raise KitError('tipografia deve ser "office" ou "editorial"')
        self.tipografia = tipografia
        self.serif = kits.TIPOGRAFIA[tipografia]["serif"]
        self.sans = kits.TIPOGRAFIA[tipografia]["sans"]
        self.doc = Document()
        self.titulo_doc = titulo
        self.cliente = cliente
        self.emissor = emissor
        self.subtitulo = subtitulo
        self.tipo = tipo if tipo is not None else self.regras["tipo"]
        self.confidencial = bool(confidencial)
        self.idioma = idioma
        self.data_str = data_str or fmt.data()
        self.pal = paleta_para(cor_marca)
        # Overrides explícitos do preset (None = decide o preset).
        self._quer_capa = capa
        self._quer_sumario = sumario
        self._quer_contracapa = contracapa
        self._capa_feita = False
        self._sumario_feito = False
        self._contracapa_feita = False
        self._sec = 0     # numeração automática de seção
        self._sub = 0     # numeração da subseção no preset "parecer"
        self._tab = 0     # rótulo "Tabela N"
        self._fig = 0     # rótulo "Gráfico N"/"Figura N"
        self._titulos = []            # [(nível, texto rotulado)] para o sumário
        self._entradas_sumario = []   # runs do número de página, para o 2º passo
        self._titulo_sumario = ""     # título do bloco de sumário, achado no PDF
        self._kpis_escritos = []      # valores exibidos, para a auditoria
        self._assinantes = []         # (nome, cargo), para achar assinatura órfã
        self._textos_de_fechamento = []   # fecho, contatos, nota
        self._contracapa_de_pagina = False
        self._fecho_pendente = None
        self._achados_geracao = []    # erros detectados durante a montagem
        self._temporarios = []        # PNGs de gráfico, apagados no fim
        st = self.doc.styles["Normal"]
        _fonte_do_estilo(st, self.sans)
        st.font.size = Pt(ESCALA["corpo"])
        st.font.color.rgb = RGBColor.from_string(self.pal["corpo"])
        _idioma(st, idioma)
        self._preparar_estilos_de_titulo()
        for s in self.doc.sections:
            s.top_margin = s.bottom_margin = Cm(margem_cm)
            s.left_margin = s.right_margin = Cm(margem_cm)
        #: Largura útil em cm — usada pelos blocos de largura fixa.
        self._larg = 21.0 - 2 * margem_cm
        self._rodape()

    # ---------- estilos de título REAIS do Word ----------
    def _preparar_estilos_de_titulo(self):
        """Dá aos estilos "Heading 1/2/3" a aparência do kit.

        Usar o estilo de verdade (em vez de um parágrafo formatado à mão) é o
        que devolve o painel de navegação no Word e os marcadores (outline) no
        PDF gêmeo — e nenhum dos dois existia enquanto o título era só um
        parágrafo em negrito."""
        tamanhos = {"Heading 1": ESCALA["h1"], "Heading 2": ESCALA["h2"],
                    "Heading 3": ESCALA["h3"]}
        for nome, tamanho in tamanhos.items():
            try:
                estilo = self.doc.styles[nome]
            except KeyError:                                   # pragma: no cover
                continue
            _fonte_do_estilo(estilo, self.serif)
            estilo.font.size = Pt(tamanho)
            estilo.font.bold = True
            estilo.font.color.rgb = RGBColor.from_string(self.pal["tinta"])
            _idioma(estilo, self.idioma)
            pf = estilo.paragraph_format
            pf.keep_with_next = True
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
            pf.left_indent = Pt(RECUO_PT)

    # ---------- primitivas de grade ----------
    def _fit(self, t, pct=5000):
        """Faz a tabela ocupar a largura útil da página (100% = pct 5000). Sem
        isso o python-docx cria a tabela com largura "auto", e conteúdo largo
        empurra as colunas PARA FORA da margem direita — o clássico "tabela
        vazando". Com tblW=100% + layout autofit, o Word ajusta as colunas ao
        conteúdo mas mantém a tabela dentro da área de texto."""
        tblPr = t._tbl.tblPr
        _tblpr_put(tblPr, _el("w:tblW", w=pct, type="pct"))
        _tblpr_put(tblPr, _el("w:tblLayout", type="autofit"))
        t.autofit = True
        return t

    def _fixa(self, t, larguras_cm):
        """Layout FIXO com larguras exatas (KPIs, linha do tempo, assinaturas).

        A largura total vai declarada em `w:tblW type="dxa"` E a grade
        (`w:tblGrid`) recebe as mesmas medidas. Sem mexer na grade, o Word
        resolve o layout fixo pela grade herdada do python-docx — colunas
        iguais, somando MAIS que a área útil — e a tabela vaza a margem mesmo
        com cada célula medida uma a uma."""
        larguras = [Cm(w).twips for w in larguras_cm]
        excesso = sum(larguras) - Cm(self._larg).twips
        if excesso > 0 and larguras:
            larguras[larguras.index(max(larguras))] -= excesso
        tblPr = t._tbl.tblPr
        _tblpr_put(tblPr, _el("w:tblW", w=sum(larguras), type="dxa"))
        _tblpr_put(tblPr, _el("w:tblLayout", type="fixed"))
        t.autofit = False
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            for col, w in zip(grid.findall(qn("w:gridCol")), larguras):
                col.set(qn("w:w"), str(w))
        for row in t.rows:
            for i, w in enumerate(larguras):
                if i < len(row.cells):
                    row.cells[i].width = Twips(w)
        return t

    def _cell_borders(self, cell, **lados):
        """lados: top/bottom/left/right = (espessura, cor) | "nil"."""
        tcPr = cell._tc.get_or_add_tcPr()
        bd = _el("w:tcBorders")
        for lado, spec in lados.items():
            if spec == "nil":
                bd.append(_el("w:" + lado, val="nil"))
            else:
                sz, cor = spec
                bd.append(_el("w:" + lado, val="single", sz=sz, space="0", color=cor))
        tcPr.append(bd)
        return cell

    def _sem_bordas(self, t):
        b = _el("w:tblBorders")
        for e in ("top", "bottom", "left", "right", "insideH", "insideV"):
            b.append(_el("w:" + e, val="nil"))
        _tblpr_put(t._tbl.tblPr, b)
        return t

    def _rotulo_bloco(self, prefixo, contador, texto):
        """Rótulo "Tabela 1 — ..." / "Gráfico 1 — ...", preso ao bloco seguinte
        por keep_with_next (solto, ele fica órfão no pé da página)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        p.paragraph_format.keep_with_next = True
        _run(p, "%s %d — %s" % (prefixo, contador, texto), bold=True,
             size=ESCALA["pequeno"], cor=self.pal["tinta"], serif=True)
        return p

    def _gap(self, cm_):
        """Espaço vertical de altura EXATA (para a capa não depender de contar
        parágrafos em branco, que estouram para uma 2ª página quando o título
        quebra em duas linhas)."""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Cm(cm_)
        pf.line_spacing = Pt(2)
        r = p.add_run(); r.font.size = Pt(2)
        return p

    def _barra(self, cor=None, tamanho=48):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(10)
        pbdr = _el("w:pBdr")
        pbdr.append(_el("w:bottom", val="single", sz=tamanho, space="1",
                        color=cor or self.pal["primaria"]))
        p._p.get_or_add_pPr().append(pbdr)
        return p

    def _mover_para_o_inicio(self, desde):
        """Move para o TOPO do corpo os elementos criados a partir do índice
        `desde`. É o que permite gerar capa e sumário no fim (quando já se sabe
        quantas seções o documento tem e em que página cada uma caiu) e ainda
        assim entregá-los na frente."""
        body = self.doc.element.body
        novos = _elementos_do_corpo(self.doc)[desde:]
        for el in reversed(novos):
            body.insert(0, el)
        return novos

    # ---------- capa e fechamento ----------
    def capa(self, estilo=None):
        """Capa em página própria. `estilo="faixa"` (padrão dos presets
        gerencial/proposta) põe o título numa faixa de tinta no terço superior;
        `estilo="simples"` mantém a capa branca com filete, para parecer e laudo.

        A capa e o fechamento têm o MESMO peso visual de propósito: capa branca
        com contracapa verde de página inteira abre discreto e fecha gritando."""
        estilo = estilo or self.regras["capa"] or "simples"
        # A capa (1ª página) não leva o rodapé "Página X de Y".
        self.doc.sections[0].different_first_page_header_footer = True
        if estilo == "faixa":
            self._capa_faixa()
        else:
            self._capa_simples()
        self._meta_da_capa()
        self.doc.add_page_break()
        self._capa_feita = True
        return self

    def _capa_faixa(self):
        t = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(t)
        cell = t.rows[0].cells[0]
        sombrear_celula(cell, self.pal["tinta"])
        self._cell_borders(cell, bottom=(24, self.pal["latao"]))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        _cell_pad(cell, 300, 300, RECUO_DXA + 120, RECUO_DXA)
        _altura_linha(t.rows[0], 9.0)
        _nao_quebrar_linha(t.rows[0])
        _limpa_paragrafo(cell.paragraphs[0])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(14)
        if self.emissor:
            _run(p, self.emissor, size=ESCALA["kicker"], cor=self.pal["latao_claro"],
                 bold=True, caps=True, spacing=44)
        if self.tipo:
            pk = cell.add_paragraph()
            pk.paragraph_format.space_after = Pt(4)
            _run(pk, self.tipo, size=ESCALA["capa_tipo"], cor=self.pal["latao_claro"],
                 bold=True, caps=True, spacing=44)
        pt_ = cell.add_paragraph()
        pt_.paragraph_format.space_after = Pt(4)
        _run(pt_, self.titulo_doc, size=ESCALA["capa_titulo"], branco=True,
             bold=True, serif=True)
        if self.subtitulo:
            ps = cell.add_paragraph()
            _run(ps, self.subtitulo, size=ESCALA["capa_sub"], cor=self.pal["creme"],
                 serif=True, italico=True)
        self._fit(t)
        # Metade inferior branca: o bloco de identificação não fica colado na
        # faixa, senão a capa vira um cabeçalho com um rodapé de 12 cm em branco.
        self._gap(6.4)

    def _capa_simples(self):
        if self.emissor:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, self.emissor, size=ESCALA["capa_meta"], cor=self.pal["cinza"],
                 bold=True, caps=True, spacing=40)
        self._gap(5.2)
        self._barra(tamanho=40)
        if self.tipo:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, self.tipo, size=ESCALA["capa_tipo"], cor=self.pal["latao"],
                 bold=True, caps=True, spacing=44)
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        p.paragraph_format.keep_with_next = True
        _run(p, self.titulo_doc, size=ESCALA["capa_titulo"], cor=self.pal["tinta"],
             bold=True, serif=True)
        if self.subtitulo:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, self.subtitulo, size=ESCALA["capa_sub"], cor=self.pal["cinza"],
                 serif=True, italico=True)
        self._barra(cor=self.pal["latao"], tamanho=12)
        self._gap(3.4)

    def _meta_da_capa(self):
        for rot, val in (("Cliente", self.cliente), ("Data", self.data_str),
                         ("Emitido por", self.emissor)):
            if val:
                p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Pt(RECUO_PT)
                _run(p, rot + ": ", bold=True, size=ESCALA["capa_meta"], cor=self.pal["cinza"])
                _run(p, val, size=ESCALA["capa_meta"], cor=self.pal["cinza"])
        if self.confidencial:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, "CONFIDENCIAL", bold=True, size=ESCALA["kicker"],
                 cor=self.pal["latao"], spacing=48)

    # ---------- títulos ----------
    def _rotulo_de_secao(self, nivel):
        """Kicker/numeração conforme o preset: "SEÇÃO 01" no gerencial, "1."/"1.1"
        no parecer, nada na proposta e na carta."""
        numeracao = self.regras["numeracao"]
        if numeracao == "secao":
            if nivel == 1:
                self._sec += 1
                self._sub = 0
                return "SEÇÃO %02d" % self._sec, None
            return None, None
        if numeracao == "decimal":
            if nivel == 1:
                self._sec += 1
                self._sub = 0
                return None, "%d." % self._sec
            self._sub += 1
            return None, "%d.%d" % (self._sec, self._sub)
        return None, None

    def titulo(self, texto, nivel=1, kicker=None):
        """Título de seção com o estilo de título REAL do Word (painel de
        navegação + marcadores no PDF).

        A numeração é do KIT: no preset gerencial sai "SEÇÃO 01" acima do
        título; no parecer, "1." / "1.1" antes dele. Nunca escreva o número no
        texto. `kicker=""` desliga; `kicker="ANEXO A"` substitui."""
        nivel = 1 if nivel not in (1, 2, 3) else nivel
        automatico, prefixo = self._rotulo_de_secao(nivel)
        if kicker == "":
            automatico = None
        elif kicker:
            automatico = kicker
        if automatico:
            k = self.doc.add_paragraph()
            k.paragraph_format.space_before = Pt(16)
            k.paragraph_format.space_after = Pt(2)
            k.paragraph_format.left_indent = Pt(RECUO_PT)
            k.paragraph_format.keep_with_next = True
            _run(k, automatico, bold=True, size=ESCALA["kicker"],
                 cor=self.pal["latao"], spacing=44)
        rotulado = ("%s %s" % (prefixo, texto)) if prefixo else str(texto)
        p = self.doc.add_paragraph(style="Heading %d" % nivel)
        p.paragraph_format.space_before = Pt(0 if automatico else (16 if nivel == 1 else 10))
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        # Recuo EXPLÍCITO além do que o estilo já traz: a aresta esquerda única
        # não pode depender de um estilo que o modelo do Word do cliente
        # eventualmente sobrescreva.
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        _run(p, rotulado, bold=True,
             size=ESCALA["h1"] if nivel == 1 else (ESCALA["h2"] if nivel == 2 else ESCALA["h3"]),
             cor=self.pal["tinta"], serif=True)
        if nivel == 1:
            pbdr = _el("w:pBdr")
            pbdr.append(_el("w:left", val="single", sz=24, space="8",
                            color=self.pal["tinta"]))
            p._p.get_or_add_pPr().append(pbdr)
        self._titulos.append((nivel, rotulado))
        return p

    # ---------- texto ----------
    def paragrafo(self, texto, justificado=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = kits.ENTRELINHA
        # O corpo usa o MESMO recuo do título e da primeira coluna da tabela:
        # é o que mantém uma única aresta esquerda na página inteira.
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        if justificado is None:
            justificado = self.regras["corpo"] == "justificado"
        if justificado:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _run(p, texto)
        return p

    def lista(self, itens, ordenada=False):
        """Lista com marcador e recuo pendurado, alinhada à mesma aresta.

        Lista curta (até 6 itens) é indivisível: um item sozinho no topo da
        página seguinte lê como falha de diagramação."""
        itens = list(itens)
        saida = []
        curta = len(itens) <= 6
        for i, item in enumerate(itens, start=1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(3)
            pf.line_spacing = kits.ENTRELINHA
            pf.left_indent = Pt(RECUO_PT + 14)
            pf.first_line_indent = Pt(-14)   # marcador pendurado em RECUO_PT
            if curta and i < len(itens):
                pf.keep_with_next = True
            _run(p, ("%d. " % i) if ordenada else "• ")
            _run(p, item)
            saida.append(p)
        return saida

    # ---------- tabela ----------
    def _bordas_tabela(self, t):
        b = _el("w:tblBorders")
        b.append(_el("w:top", val="single", sz=4, color=self.pal["borda"]))
        b.append(_el("w:bottom", val="single", sz=4, color=self.pal["borda"]))
        b.append(_el("w:insideH", val="single", sz=4, color=self.pal["borda"]))
        b.append(_el("w:left", val="nil"))
        b.append(_el("w:right", val="nil"))
        b.append(_el("w:insideV", val="nil"))
        _tblpr_put(t._tbl.tblPr, b)

    def tabela(self, cabecalho, linhas, total=False, titulo=None, fonte=None,
               moeda=None, pct=None, milhar=None, data=None, larguras=None):
        """Tabela do kit — **passe NÚMEROS, não strings formatadas**.

        `moeda=`/`pct=`/`milhar=`/`data=` recebem NOMES de coluna do cabeçalho;
        o kit formata em pt-BR, alinha o numérico à direita e escreve o negativo
        como "(1.234,56)" em vermelho, a convenção do demonstrativo contábil.

        `total="soma"` CALCULA a linha de total das colunas numéricas (é o que
        impede um total que não fecha com o corpo da tabela). `total=True`
        mantém o comportamento antigo: a última linha que você passou já é o
        total e só recebe o destaque. `titulo=` põe o rótulo "Tabela N — ..."
        acima e `fonte=` a linha de origem do dado abaixo, colada na tabela."""
        cabecalho = list(cabecalho)
        tipos = tipos_de_coluna(cabecalho, moeda, pct, milhar, data)
        linhas, achados = normaliza_linhas(cabecalho, linhas)
        self._achados_geracao.extend(achados)
        if str(total).lower() == "soma":
            linhas = linhas + [linha_de_total(cabecalho, linhas, tipos)]
            destaca_ultima = True
        else:
            destaca_ultima = bool(total)
        if titulo:
            self._tab += 1
            self._rotulo_bloco("Tabela", self._tab, titulo)
        t = self.doc.add_table(rows=1, cols=len(cabecalho))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            t.style = "Table Grid"
        except Exception:                                      # pragma: no cover
            pass
        self._bordas_tabela(t)
        # cabeçalho (repetido em quebra de página)
        hdr = t.rows[0]
        hdr._tr.get_or_add_trPr().append(_el("w:tblHeader"))
        for i, c in enumerate(cabecalho):
            cell = hdr.cells[i]
            sombrear_celula(cell, self.pal["tinta"])
            _limpa_paragrafo(cell.paragraphs[0])
            _run(cell.paragraphs[0], c, bold=True, branco=True, size=ESCALA["corpo"])
            if i in tipos and tipos[i] != "data":
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # A 1ª coluna recebe o recuo da grade; as demais, só o respiro.
            _cell_pad(cell, l=RECUO_DXA if i == 0 else RESPIRO_DXA)
        # linhas
        n = len(linhas)
        for ri, linha in enumerate(linhas):
            row = t.add_row()
            eh_total = destaca_ultima and ri == n - 1
            zebra = (not eh_total) and n >= 6 and (ri % 2 == 1)
            for i, v in enumerate(linha):
                cell = row.cells[i]
                if eh_total or zebra:
                    sombrear_celula(cell, self.pal["suave"])
                negativo = e_numero(v) and float(v) < 0 and tipos.get(i) in ("moeda", "milhar", "pct")
                texto = formata_valor(v, tipos.get(i), parenteses=negativo)
                _limpa_paragrafo(cell.paragraphs[0])
                cor = self.pal["negativo"] if negativo else (self.pal["tinta"] if eh_total else None)
                _run(cell.paragraphs[0], texto, bold=eh_total, size=ESCALA["corpo"], cor=cor)
                direita = e_numero(v) or (i in tipos and tipos[i] != "data") or _num_texto(texto)
                cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.RIGHT if direita
                                                else WD_ALIGN_PARAGRAPH.LEFT)
                _cell_pad(cell, l=RECUO_DXA if i == 0 else RESPIRO_DXA)
            if eh_total:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    bd = _el("w:tcBorders")
                    # Filete de latão: separa o TOTAL sem repetir a cor do
                    # cabeçalho — é o acento da identidade.
                    bd.append(_el("w:top", val="single", sz=14, color=self.pal["latao"]))
                    tcPr.append(bd)
        if larguras:
            self._fixa(t, larguras)
        else:
            self._fit(t)
        self._paginacao_tabela(t, bool(fonte))
        if fonte:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, fonte, size=ESCALA["legenda"], cor=self.pal["cinza"])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    def _paginacao_tabela(self, t, tem_fonte):
        """Regras de quebra da tabela — código, não instrução ao modelo.

        Até `LIMITE_TABELA_INDIVISIVEL` linhas a tabela inteira viaja junto
        (cabeçalho repetido com TOTAL órfão numa página só é o defeito mais
        visível de todos). Acima disso, o cabeçalho se repete e a PENÚLTIMA
        linha "mantém com a próxima", para o TOTAL nunca abrir a página
        sozinho. A linha "Fonte:" cola na última linha da tabela."""
        linhas = t.rows
        for row in linhas:
            _nao_quebrar_linha(row)
        if len(linhas) <= LIMITE_TABELA_INDIVISIVEL:
            for row in linhas[:-1]:
                _colar_na_proxima(row)
        elif len(linhas) >= 2:
            _colar_na_proxima(linhas[-2])
        if tem_fonte and linhas:
            _colar_na_proxima(linhas[-1])

    # ---------- destaques ----------
    def callout(self, rotulo, texto, tipo="info"):
        """Caixa de destaque. `tipo`: info | alerta | critico | sucesso."""
        fundos = {"info": self.pal["info_bg"], "alerta": self.pal["alerta_bg"],
                  "critico": self.pal["critico_bg"], "sucesso": self.pal["sucesso_bg"]}
        bordas = {"info": self.pal["info_bd"], "alerta": self.pal["alerta_bd"],
                  "critico": self.pal["critico_bd"], "sucesso": self.pal["sucesso_bd"]}
        if tipo not in fundos:
            raise KitError('callout: tipo deve ser info | alerta | critico | sucesso (recebi %r)' % (tipo,))
        t = self.doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        sombrear_celula(cell, fundos[tipo])
        _cell_pad(cell, 120, 120, RECUO_DXA, RECUO_DXA)
        tcPr = cell._tc.get_or_add_tcPr()
        b = _el("w:tcBorders")
        b.append(_el("w:left", val="single", sz=24, color=bordas[tipo]))
        for e in ("top", "bottom", "right"):
            b.append(_el("w:" + e, val="nil"))
        tcPr.append(b)
        _limpa_paragrafo(cell.paragraphs[0])
        if rotulo:
            _run(cell.paragraphs[0], rotulo, bold=True, size=ESCALA["pequeno"],
                 cor=bordas[tipo], caps=True, spacing=20)
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        _run(p, texto)
        self._fit(t)
        _nao_quebrar_linha(t.rows[0])     # callout é indivisível
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    def kpis(self, itens):
        """Cartões de indicador — de 2 a 6.

        Cada item é `(valor, rótulo)` ou `(valor, rótulo, tipo)` com tipo
        `moeda | pct | milhar | num | texto`. O valor pode ser NÚMERO (o kit
        formata) ou texto já pronto. O tamanho do número é adaptativo: valor
        comprido reduz de corpo em vez de quebrar em duas linhas — era isso que
        deixava "R$ 1,89 / mi" ao lado de "32,5%" com alturas diferentes."""
        itens = list(itens)
        if not itens:
            return None
        if len(itens) > 6:
            raise KitError("kpis: use de 2 a 6 indicadores (recebi %d)" % len(itens))
        ultima = None
        for fileira in linhas_de_kpi(itens):
            ultima = self._fileira_de_kpis(fileira)
        return ultima

    def _fileira_de_kpis(self, fileira):
        textos = []
        for item in fileira:
            item = list(item)
            valor, rotulo = item[0], item[1]
            tipo = item[2] if len(item) > 2 else None
            tipo = None if tipo in (None, "texto", "num") else tipo
            textos.append((formata_valor(valor, tipo), rotulo))
        # Um só tamanho para a fileira inteira: cartões vizinhos com corpos
        # diferentes leem como erro, não como hierarquia.
        tamanho = min(escala_kpi(v) for v, _ in textos)
        t = self.doc.add_table(rows=2, cols=len(textos))
        for i, (valor, rot) in enumerate(textos):
            self._kpis_escritos.append(valor)
            cv = t.rows[0].cells[i]
            sombrear_celula(cv, self.pal["suave"]); _cell_pad(cv, 130, 40, 120, 120)
            _limpa_paragrafo(cv.paragraphs[0])
            cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(cv.paragraphs[0], valor or "—", bold=True, size=tamanho,
                 cor=self.pal["tinta"], serif=True)
            # Filete de latão no topo do cartão — o acento da identidade.
            self._cell_borders(cv, top=(24, self.pal["latao"]))
            cr = t.rows[1].cells[i]
            sombrear_celula(cr, self.pal["suave"]); _cell_pad(cr, 0, 130, 120, 120)
            _limpa_paragrafo(cr.paragraphs[0])
            cr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(cr.paragraphs[0], rot, size=ESCALA["kpi_rotulo"], bold=True,
                 cor=self.pal["cinza"], caps=True, spacing=24)
        # Altura mínima igual nas duas linhas: é o que iguala os cartões mesmo
        # quando um rótulo quebra em duas linhas e o vizinho não.
        _altura_linha(t.rows[0], 1.15, exata=False)
        _altura_linha(t.rows[1], 0.75, exata=False)
        for row in t.rows:
            _nao_quebrar_linha(row)
        _colar_na_proxima(t.rows[0])
        # Separa os cartões: linha branca grossa entre colunas (senão os KPIs
        # colados com o mesmo fundo suave viram um bloco único, sem respiro).
        b = _el("w:tblBorders")
        for e in ("top", "bottom", "left", "right", "insideH"):
            b.append(_el("w:" + e, val="nil"))
        b.append(_el("w:insideV", val="single", sz=36, color=self.pal["branco"]))
        _tblpr_put(t._tbl.tblPr, b)
        self._fit(t)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    # ---------- blocos editoriais ----------
    def sumario(self, entradas=None, titulo="Conteúdo deste documento"):
        """Sumário com índice pontilhado.

        Sem argumento, monta-se dos `titulo()` já registrados e as páginas são
        as REAIS: `salvar()` grava, converte para PDF, descobre em que página
        cada título caiu e reescreve os números antes de converter de novo.
        `sumario(entradas)` continua aceito para forçar valores à mão — com
        aviso, porque foi assim que o sumário do relatório saiu apontando a
        página 3 para uma seção que estava na 2."""
        if entradas is not None:
            warnings.warn(
                "docpro: sumario(entradas) informa as páginas à mão e elas não "
                "são conferidas; prefira sumario() sem argumento — o kit "
                "descobre a página real no PDF.")
        k = self.doc.add_paragraph()
        k.paragraph_format.space_before = Pt(6)
        k.paragraph_format.space_after = Pt(2)
        k.paragraph_format.left_indent = Pt(RECUO_PT)
        k.paragraph_format.keep_with_next = True
        _run(k, "SUMÁRIO", bold=True, size=ESCALA["kicker"], cor=self.pal["latao"],
             spacing=44)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        p.paragraph_format.keep_with_next = True
        pbdr = _el("w:pBdr")
        pbdr.append(_el("w:bottom", val="single", sz=20, space="8", color=self.pal["tinta"]))
        p._p.get_or_add_pPr().append(pbdr)
        _run(p, titulo, bold=True, size=ESCALA["h1"] + 2, cor=self.pal["tinta"], serif=True)
        self._titulo_sumario = str(titulo)
        if entradas is None:
            entradas = [(t, "") for nivel, t in self._titulos if nivel == 1]
        for i, entrada in enumerate(entradas):
            entrada = list(entrada)
            if len(entrada) == 3:
                num, tit, pag = entrada
            else:
                tit, pag = (entrada + [""])[:2]
                num = "%02d" % (i + 1)
            e = self.doc.add_paragraph()
            e.paragraph_format.space_before = Pt(6)
            e.paragraph_format.space_after = Pt(6)
            e.paragraph_format.left_indent = Pt(RECUO_PT)
            e.paragraph_format.tab_stops.add_tab_stop(
                Cm(self._larg), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            pbdr = _el("w:pBdr")
            pbdr.append(_el("w:bottom", val="single", sz=4, space="4",
                            color=self.pal["borda_leve"]))
            e._p.get_or_add_pPr().append(pbdr)
            _run(e, str(num) + "   ", bold=True, size=ESCALA["apoio"], cor=self.pal["latao"])
            _run(e, tit, bold=True, size=ESCALA["corpo"] + 0.5, serif=True)
            e.add_run("\t")
            run_pagina = _run(e, str(pag), size=ESCALA["apoio"], cor=self.pal["cinza"])
            # Guarda o run para o 2º passo: é nele que a página REAL é escrita.
            self._entradas_sumario.append((tit, run_pagina))
        self.doc.add_paragraph().paragraph_format.space_after = Pt(8)
        self._sumario_feito = True
        return None

    def citacao(self, texto, fonte=""):
        """Frase-chave destacada por um filete de latão à esquerda."""
        t = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(t)
        cell = t.rows[0].cells[0]
        self._cell_borders(cell, left=(18, self.pal["latao"]))
        _cell_pad(cell, 60, 60, RECUO_DXA + 180, RECUO_DXA)
        _limpa_paragrafo(cell.paragraphs[0])
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.35
        _run(p, "“" + str(texto).strip("“”\"") + "”", size=ESCALA["h2"] + 1,
             cor=self.pal["tinta"], serif=True, italico=True)
        if fonte:
            pf = cell.add_paragraph()
            pf.paragraph_format.space_before = Pt(5)
            _run(pf, fonte, bold=True, size=ESCALA["legenda"], cor=self.pal["cinza"],
                 caps=True, spacing=20)
        self._fit(t)
        _nao_quebrar_linha(t.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    def etapas(self, itens, titulo=None):
        """Linha do tempo vertical — `[(etapa, quando/descrição), ...]`."""
        itens = list(itens)
        if not itens:
            return None
        if titulo:
            self.titulo(titulo, nivel=2)
        col_num = 1.3
        t = self.doc.add_table(rows=len(itens), cols=2)
        self._sem_bordas(t)
        for i, item in enumerate(itens):
            etapa, quando = (list(item) + [""])[:2]
            cnum = t.rows[i].cells[0]
            _limpa_paragrafo(cnum.paragraphs[0])
            _run(cnum.paragraphs[0], "%02d" % (i + 1), bold=True, size=ESCALA["h2"] + 1.5,
                 cor=self.pal["latao"], serif=True)
            _cell_pad(cnum, 110, 110, RECUO_DXA, 60)
            ctx = t.rows[i].cells[1]
            _limpa_paragrafo(ctx.paragraphs[0])
            _run(ctx.paragraphs[0], etapa, bold=True, size=ESCALA["corpo"])
            if quando:
                pd = ctx.add_paragraph()
                pd.paragraph_format.space_before = Pt(1)
                _run(pd, quando, size=ESCALA["apoio"], cor=self.pal["cinza"])
            _cell_pad(ctx, 110, 110, 60, RECUO_DXA)
            if i < len(itens) - 1:
                self._cell_borders(cnum, bottom=(4, self.pal["borda_leve"]))
                self._cell_borders(ctx, bottom=(4, self.pal["borda_leve"]))
            _nao_quebrar_linha(t.rows[i])
            # Cronograma curto é indivisível: uma etapa sozinha no topo da
            # página seguinte lê como falha de diagramação, não como conteúdo.
            if len(itens) <= 6 and i < len(itens) - 1:
                _colar_na_proxima(t.rows[i])
        self._fixa(t, [col_num, self._larg - col_num])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return t

    def imagem(self, caminho, largura_cm=None, legenda=""):
        """Imagem contida na largura útil, com o rótulo "Figura N"."""
        if not os.path.exists(caminho):
            raise KitError("imagem: arquivo não encontrado: %s" % caminho)
        if legenda:
            self._fig += 1
            self._rotulo_bloco("Figura", self._fig, legenda)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        p.add_run().add_picture(caminho, width=Cm(min(largura_cm or self._larg, self._larg)))
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return p

    # ---------- gráficos ----------
    def _grafico(self, gerar, titulo, largura_cm, altura_cm, *args, **kwargs):
        self._fig += 1
        caminho = os.path.join(tempfile.gettempdir(),
                               "docpro_fig_%d_%d.png" % (id(self) % 100000, self._fig))
        # Largura CHEIA: um gráfico com 60% da largura ao lado de uma tabela com
        # 100% cria duas manchas de larguras diferentes na mesma página.
        gerar(caminho, *args, pal=self.pal, largura_cm=largura_cm or self._larg,
              altura_cm=altura_cm, **kwargs)
        self._temporarios.append(caminho)
        if titulo:
            self._rotulo_bloco("Gráfico", self._fig, titulo)
        # O parágrafo da figura recebe o MESMO recuo do corpo: sem ele a imagem
        # nasce colada na margem e cria uma segunda aresta na página.
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(RECUO_PT)
        p.add_run().add_picture(caminho, width=Cm(self._larg - 2 * RECUO_PT / 28.35))
        self.doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return caminho

    def grafico_barras(self, categorias, series, titulo="", largura_cm=None,
                       altura_cm=7.0, sufixo_eixo=""):
        """`series` = {"Nome": [valores]} ou uma lista simples de valores."""
        return self._grafico(kits.grafico_barras_png, titulo, largura_cm, altura_cm,
                             categorias, series, sufixo_eixo=sufixo_eixo)

    def grafico_linhas(self, categorias, series, titulo="", largura_cm=None,
                       altura_cm=7.0, sufixo_eixo=""):
        return self._grafico(kits.grafico_linhas_png, titulo, largura_cm, altura_cm,
                             categorias, series, sufixo_eixo=sufixo_eixo)

    def grafico_pizza(self, rotulos, valores, titulo="", largura_cm=11.0,
                      altura_cm=7.0):
        return self._grafico(kits.grafico_pizza_png, titulo, largura_cm, altura_cm,
                             rotulos, valores)

    # ---------- fecho, assinaturas e fechamento ----------
    def fecho(self, local_data):
        """Local e data. Guardado para entrar no MESMO bloco indivisível das
        assinaturas — o fecho numa página e a assinatura na seguinte é o defeito
        que mais desmoraliza um documento pronto."""
        self._fecho_pendente = str(local_data)
        return self

    def assinaturas(self, nomes, cargos=None, local_data=None, subtitulos=None,
                    testemunhas=None):
        """Linhas de assinatura em PARES lado a lado, dentro da grade.

        O fecho, as assinaturas e as testemunhas vão numa tabela de UMA coluna
        que não se parte: nunca sobra uma página só com a linha de assinatura.
        `subtitulos=` é aceito como alias de `cargos=` (compatibilidade)."""
        nomes = list(nomes)
        cargos = list(cargos or subtitulos or []) + [None] * len(nomes)
        if local_data:
            self.fecho(local_data)
        pendente = self._fecho_pendente
        self._fecho_pendente = None
        self._assinantes += [(n, cargos[i]) for i, n in enumerate(nomes)]
        self._assinantes += [(n, None) for n in list(testemunhas or [])]
        if pendente:
            self._textos_de_fechamento.append(pendente)

        externa = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(externa)
        _nao_quebrar_linha(externa.rows[0])
        celula = externa.rows[0].cells[0]
        _cell_pad(celula, 0, 0, 0, 0)
        _limpa_paragrafo(celula.paragraphs[0])
        if pendente:
            p = celula.paragraphs[0]
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, pendente)
        self._bloco_de_assinaturas(celula, nomes, cargos)
        if testemunhas:
            pt_ = celula.add_paragraph()
            pt_.paragraph_format.space_before = Pt(18)
            pt_.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(pt_, "TESTEMUNHAS", bold=True, size=ESCALA["kicker"],
                 cor=self.pal["cinza"], caps=True, spacing=40)
            self._bloco_de_assinaturas(celula, list(testemunhas),
                                       [None] * len(list(testemunhas)))
        self._fit(externa)
        return externa

    def _bloco_de_assinaturas(self, celula, nomes, cargos):
        """Escreve pares de assinatura DENTRO de uma célula (o bloco
        indivisível). Uma tabela aninhada por par mantém o alinhamento."""
        meio = 1.4
        w = (self._larg - meio) / 2
        for pi in range(0, len(nomes), 2):
            par = nomes[pi:pi + 2]
            espaco = celula.add_paragraph()
            espaco.paragraph_format.space_after = Cm(1.8 if pi == 0 else 1.2)
            espaco.paragraph_format.line_spacing = Pt(2)
            espaco.add_run().font.size = Pt(2)
            t = celula.add_table(rows=1, cols=3)
            self._sem_bordas(t)
            _nao_quebrar_linha(t.rows[0])
            # Assinatura SOLITÁRIA vai ao centro: encostada à esquerda, num par
            # que não existe, ela lê como se faltasse o segundo signatário.
            sozinho = len(par) == 1
            colunas = [(self._larg - w) / 2, w, (self._larg - w) / 2] if sozinho \
                else [w, meio, w]
            for j, nome in enumerate(par):
                cell = t.rows[0].cells[1 if sozinho else j * 2]
                self._cell_borders(cell, top=(6, self.pal["corpo"]))
                _cell_pad(cell, 110, 40, 60, 60)
                _limpa_paragrafo(cell.paragraphs[0])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(cell.paragraphs[0], nome, bold=True, size=ESCALA["corpo"] + 0.5,
                     serif=True)
                cargo = cargos[pi + j] if pi + j < len(cargos) else None
                if cargo:
                    pc = cell.add_paragraph()
                    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(pc, liga_identificadores(cargo), size=ESCALA["legenda"] + 0.5,
                         cor=self.pal["cinza"])
            self._fixa(t, colunas)

    def contracapa(self, contatos=None, nota=None, estilo=None):
        """Fechamento: faixa de tinta no fim da última página (padrão) ou página
        inteira (`estilo="pagina"`).

        A página inteira só se justifica num documento longo: num relatório de
        cinco páginas ela vira uma sexta página verde e vazia, e o documento
        fecha gritando o que a capa abriu discreto."""
        estilo = estilo or self._quer_contracapa or "faixa"
        if estilo == "pagina":
            self._contracapa_de_pagina = True
            self.doc.add_page_break()
        self._faixa_de_fechamento(contatos, nota, alta=(estilo == "pagina"))
        self._contracapa_feita = True
        return self

    def _ancorar_no_pe(self, t):
        """Ancora a tabela no PÉ da mancha da página (acima do rodapé).

        Sem isto a faixa de fechamento nasce logo depois da assinatura e deixa
        meia página em branco abaixo dela — o oposto de "fechamento". Com
        `vertAnchor="margin"` + `tblpYSpec="bottom"` ela desce até o fim da área
        de texto e o rodapé paginado continua visível embaixo."""
        pos = _el("w:tblpPr", leftFromText=0, rightFromText=0, topFromText=180,
                  bottomFromText=0, vertAnchor="margin", horzAnchor="margin",
                  tblpXSpec="center", tblpYSpec="bottom")
        _tblpr_put(t._tbl.tblPr, pos)
        _tblpr_put(t._tbl.tblPr, _el("w:tblOverlap", val="never"))
        return t

    def _faixa_de_fechamento(self, contatos, nota, alta=False):
        self._textos_de_fechamento += [self.emissor or self.titulo_doc, nota or ""]
        self._textos_de_fechamento += list(contatos or [])
        t = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(t)
        cell = t.rows[0].cells[0]
        sombrear_celula(cell, self.pal["tinta"])
        self._cell_borders(cell, top=(24, self.pal["latao"]))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _altura_linha(t.rows[0], 23.0 if alta else 4.2)
        _nao_quebrar_linha(t.rows[0])
        _cell_pad(cell, 200, 200, RECUO_DXA, RECUO_DXA)
        _limpa_paragrafo(cell.paragraphs[0])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, self.emissor or self.titulo_doc, bold=True,
             size=20 if alta else ESCALA["h2"], cor=self.pal["suave"], serif=True)
        if contatos:
            pk = cell.add_paragraph()
            pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pk.paragraph_format.space_before = Pt(10 if alta else 4)
            _run(pk, "CONTATO", bold=True, size=ESCALA["legenda"],
                 cor=self.pal["latao_claro"], spacing=48)
            for linha in contatos:
                pc = cell.add_paragraph()
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pc.paragraph_format.space_after = Pt(2)
                _run(pc, linha, size=ESCALA["apoio"], cor=self.pal["creme"])
        if nota is None and self.confidencial and alta:
            nota = ("Este documento é confidencial e destina-se exclusivamente "
                    "ao cliente identificado na capa.")
        if nota:
            pn = cell.add_paragraph()
            pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pn.paragraph_format.space_before = Pt(36 if alta else 8)
            _run(pn, nota, size=ESCALA["legenda"], cor=self.pal["creme"])
        self._fit(t)
        if not alta:
            self._ancorar_no_pe(t)
        return t

    # ---------- rodapé e metadados ----------
    def _rodape(self):
        sec = self.doc.sections[-1]
        f = sec.footer
        f.is_linked_to_previous = False
        p = f.paragraphs[0]
        _limpa_paragrafo(p)
        pbdr = _el("w:pBdr")
        pbdr.append(_el("w:top", val="single", sz=4, color=self.pal["borda"]))
        p._p.get_or_add_pPr().append(pbdr)
        _run(p, (self.emissor or "") + "   ", size=ESCALA["rodape"], cor=self.pal["cinza"])
        if self.confidencial:
            _run(p, "CONFIDENCIAL   ", bold=True, size=ESCALA["rodape"],
                 cor=self.pal["latao"], spacing=24)
        _run(p, "Página ", size=ESCALA["rodape"], cor=self.pal["cinza"])
        self._campo(p, "PAGE")
        _run(p, " de ", size=ESCALA["rodape"], cor=self.pal["cinza"])
        self._campo(p, "NUMPAGES")

    def _campo(self, p, code):
        r = p.add_run()
        r._r.append(_el("w:fldChar", fldCharType="begin"))
        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " " + code + " "
        r2._r.append(instr)
        r3 = p.add_run()
        r3._r.append(_el("w:fldChar", fldCharType="end"))
        for rr in (r, r2, r3):
            rr.font.size = Pt(ESCALA["rodape"])
            rr.font.color.rgb = RGBColor.from_string(self.pal["cinza"])

    def _metadados(self):
        """Título, autor, assunto e idioma no .docx. Sem eles o Word mostra o
        nome do arquivo na barra e o PDF gêmeo sai sem /Title."""
        props = self.doc.core_properties
        props.title = limpa_texto(self.titulo_doc)
        props.author = limpa_texto(self.emissor or "Frederico AI Studio")
        props.subject = limpa_texto(self.tipo or self.subtitulo)
        props.keywords = limpa_texto(self.cliente)
        props.language = self.idioma
        props.category = self.preset
        props.created = props.modified = datetime.now()

    # ---------- saída e conferência ----------
    def salvar(self, caminho, pdf=True, auditar=True):
        """Grava o .docx, gera o PDF gêmeo, ACERTA o sumário com as páginas
        REAIS e AUDITA o arquivo pronto.

        Devolve `{"ok", "paginas", "achados", "pdf"}`. Achado grave levanta
        `KitError` apontando o bloco: é melhor falhar na geração do que entregar
        um documento com placeholder ou com o sumário mentindo a página."""
        self._finalizar_estrutura()
        self._metadados()
        self.doc.save(caminho)
        caminho_pdf = _converter_pdf(caminho) if pdf else None
        paginas = _texto_por_pagina(caminho_pdf) if caminho_pdf else []
        divergencias = []
        if paginas and self._entradas_sumario:
            divergencias = self._acertar_sumario(paginas)
            if divergencias:
                # 2º passo: o sumário agora conhece a página real de cada seção.
                self.doc.save(caminho)
                caminho_pdf = _converter_pdf(caminho)
                paginas = _texto_por_pagina(caminho_pdf) if caminho_pdf else paginas
        self._limpar_temporarios()
        if not auditar:
            return relatorio([], paginas=len(paginas), pdf=caminho_pdf)
        rel = relatorio(self._auditar(caminho, caminho_pdf, paginas),
                        paginas=len(paginas), pdf=caminho_pdf)
        return falha_se_grave(rel, "docpro")

    def _finalizar_estrutura(self):
        """Aplica o preset: capa, sumário e fechamento que o modelo não pediu
        explicitamente. Capa e sumário são gerados AGORA (quando já se sabe
        quantas seções existem) e movidos para o topo do corpo."""
        secoes = sum(1 for nivel, _ in self._titulos if nivel == 1)
        quer_sumario = self._quer_sumario
        if quer_sumario is None:
            quer_sumario = (self.regras["sumario"] == "auto"
                            and secoes >= MINIMO_SECOES_PARA_SUMARIO)
        if quer_sumario and not self._sumario_feito:
            marca = len(_elementos_do_corpo(self.doc))
            self.sumario()
            if len(self._entradas_sumario) > LIMITE_SUMARIO_SEM_PAGINA:
                self.doc.add_page_break()
            self._mover_para_o_inicio(marca)

        quer_capa = self._quer_capa
        if quer_capa is None:
            quer_capa = self.regras["capa"] is not None
        if quer_capa and not self._capa_feita:
            marca = len(_elementos_do_corpo(self.doc))
            self.capa(estilo=None if quer_capa is True else quer_capa)
            self._mover_para_o_inicio(marca)

        pendente = self._fecho_pendente
        if pendente:
            # `fecho()` sem `assinaturas()` depois: escreve o local e a data.
            self._fecho_pendente = None
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.left_indent = Pt(RECUO_PT)
            _run(p, pendente)

        fechamento = self._quer_contracapa
        if fechamento is None:
            fechamento = self.regras["fechamento"]
        if fechamento in ("faixa", "pagina") and not self._contracapa_feita:
            self.contracapa(estilo=fechamento)

    def _fim_do_sumario(self, paginas):
        """Índice (0-based) da PRIMEIRA página depois do sumário.

        Sem isto a busca acha cada título na própria página do sumário — que
        lista todos eles — e o resultado é um sumário mandando o leitor para a
        página do sumário. Foi exatamente esse o defeito do relatório que
        motivou a v2."""
        if not self._entradas_sumario:
            return 0
        pagina = _pagina_do_texto(paginas, self._titulo_sumario)
        if not pagina:
            return 0
        i = pagina - 1
        # O sumário pode ocupar mais de uma página: enquanto a seguinte ainda
        # tiver linhas de guia pontilhada, ela é continuação dele.
        while i + 1 < len(paginas) and re.search(r"\.{10,}", paginas[i + 1]):
            i += 1
        return i

    def _acertar_sumario(self, paginas):
        """Escreve nos runs do sumário a página REAL de cada seção. Devolve as
        entradas que mudaram (vazio = o sumário já estava certo)."""
        mudou = []
        inicio = self._fim_do_sumario(paginas)
        for titulo_txt, run_pagina in self._entradas_sumario:
            pagina = _pagina_do_texto(paginas, titulo_txt, a_partir_de=inicio,
                                      ignorar_sumario=True)
            if not pagina:
                continue
            novo = str(pagina)
            if run_pagina.text != novo:
                run_pagina.text = novo
                mudou.append((titulo_txt, pagina))
        return mudou

    def _limpar_temporarios(self):
        for caminho in self._temporarios:
            try:
                os.remove(caminho)
            except OSError:                                    # pragma: no cover
                pass
        self._temporarios = []

    # ---------- auditoria ----------
    def _auditar(self, caminho, caminho_pdf, paginas):
        from docx import Document as _Doc
        achados = list(self._achados_geracao)
        doc = _Doc(caminho)
        achados += achados_de_placeholder(textos_do_docx(doc), "documento")
        if not doc.paragraphs and not doc.tables:
            achados.append(achado("grave", "documento-vazio",
                                  "o .docx não tem parágrafo nem tabela"))
        achados += self._auditar_celulas_vazias(doc)
        if caminho_pdf is None:
            achados.append(achado(
                "aviso", "sem-pdf-gemeo",
                "o PDF gêmeo não foi gerado (LibreOffice indisponível): sumário, "
                "KPI quebrado, título órfão e página vazia não puderam ser conferidos"))
            return achados
        achados += self._auditar_pdf(paginas)
        return achados

    def _auditar_celulas_vazias(self, doc):
        """Célula de VALOR vazia numa tabela de dados. A primeira coluna pode
        ficar vazia (é rótulo de continuação); as demais, não."""
        achados = []
        for it, t in enumerate(doc.tables, start=1):
            if len(t.columns) < 2 or len(t.rows) < 2:
                continue
            for ir, row in enumerate(t.rows[1:], start=1):
                celulas = [c.text.strip() for c in row.cells]
                if not any(celulas) or celulas[0].strip().upper().startswith("TOTAL"):
                    continue
                vazias = [j for j, v in enumerate(celulas) if j > 0 and not v]
                if vazias and len(vazias) < len(celulas) - 1:
                    achados.append(achado(
                        "aviso", "celula-de-valor-vazia",
                        "tabela %d, linha %d: coluna(s) %s sem valor" % (it, ir, vazias)))
        return achados

    def _auditar_pdf(self, paginas):
        achados = []
        if not paginas:
            return [achado("aviso", "pdf-ilegivel",
                           "não foi possível ler o texto do PDF gêmeo (pypdf ausente?)")]
        # Sumário: cada entrada tem de apontar a página em que a seção está.
        inicio = self._fim_do_sumario(paginas)
        for titulo_txt, run_pagina in self._entradas_sumario:
            real = _pagina_do_texto(paginas, titulo_txt, a_partir_de=inicio,
                                    ignorar_sumario=True)
            escrita = run_pagina.text.strip()
            if real and escrita and escrita != str(real):
                achados.append(achado(
                    "grave", "sumario-divergente",
                    "o sumário diz página %s para %r, mas a seção está na página %d"
                    % (escrita, titulo_txt, real)))
        # KPI que quebrou em duas linhas: o valor não aparece inteiro em
        # nenhuma página do PDF.
        for valor in self._kpis_escritos:
            alvo = re.sub(r"\s+", " ", valor).strip()
            if alvo and not any(alvo in re.sub(r"[ \t]+", " ", p) for p in paginas):
                achados.append(achado(
                    "aviso", "kpi-quebrado",
                    "o KPI %r não aparece inteiro em nenhuma linha do PDF — "
                    "provavelmente quebrou em duas" % valor))
        # Página quase vazia e assinatura órfã — a MESMA regra dos três kits.
        achados += kits.achados_de_paginacao(
            paginas, self._assinantes, self._textos_de_fechamento,
            ultima_e_arte=self._contracapa_de_pagina)
        return achados


class Sobrio:
    """Documento SÓBRIO/registrável (ata, contrato, alteração contratual,
    distrato, procuração) em python-docx puro: ZERO cor, corpo SEMPRE
    JUSTIFICADO, fonte formal, margens de documento oficial e rodapé "Página X
    de Y". A justificação nasce no estilo Normal — não depende de o modelo
    lembrar de alinhar cada parágrafo.

    Traz os helpers de redação que a Junta Comercial espera: `clausula()`
    numera "CLÁUSULA PRIMEIRA — DO OBJETO" sozinha, `inciso()` escreve "I –",
    `paragrafo_numerado()` escreve "§ 1º" e `item()` numera "1., 2." sem
    duplicar quando o texto já vem numerado.

    Uso:
        from docpro import Sobrio
        a = Sobrio("ATA DE REUNIÃO DE SÓCIOS",
                   identificacao="ACME LTDA — CNPJ 00.000.000/0001-00")
        a.paragrafo("Aos dois dias do mês de setembro de dois mil e vinte e seis, ...")
        a.secao("ORDEM DO DIA"); a.item("Aprovação das contas de 2025;")
        a.fecho("Palmas/TO, 02 de setembro de 2026.")
        a.assinaturas(["Nome"], cargos=["Sócio-administrador"])
        rel = a.salvar("/workspace/outputs/ata.docx")
    """

    ORDINAIS = ("PRIMEIRA", "SEGUNDA", "TERCEIRA", "QUARTA", "QUINTA", "SEXTA",
                "SÉTIMA", "OITAVA", "NONA", "DÉCIMA", "DÉCIMA PRIMEIRA",
                "DÉCIMA SEGUNDA", "DÉCIMA TERCEIRA", "DÉCIMA QUARTA",
                "DÉCIMA QUINTA", "DÉCIMA SEXTA", "DÉCIMA SÉTIMA",
                "DÉCIMA OITAVA", "DÉCIMA NONA", "VIGÉSIMA")
    ROMANOS = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
               "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX")

    def __init__(self, titulo=None, identificacao=None, fonte="Times New Roman",
                 tamanho=12, entrelinha=1.5, rodape_paginado=True, rubrica=False,
                 idioma="pt-BR"):
        self.doc = Document()
        self.tam = tamanho
        self.rubrica = bool(rubrica)
        self.titulo_doc = titulo or ""
        self.idioma = idioma
        self._clausulas = 0
        self._incisos = 0
        self._paragrafos_numerados = 0
        self._itens = 0
        self._assinantes = []
        self._textos_de_fechamento = []
        # O registrável não tem contracapa: a constante existe para a
        # auditoria de paginação ser a MESMA dos dois documentos Word.
        self._contracapa_de_pagina = False
        self._achados_geracao = []
        self._fecho_pendente = None
        for s in self.doc.sections:
            s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
            s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
        st = self.doc.styles["Normal"]
        _fonte_do_estilo(st, fonte)
        st.font.size = Pt(tamanho)
        st.font.color.rgb = RGBColor(0, 0, 0)
        _idioma(st, idioma)
        pf = st.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # <- justificação de fábrica
        pf.line_spacing = entrelinha
        pf.space_after = Pt(6)
        if rodape_paginado:
            self._rodape()
        if titulo:
            self.identificacao(titulo, identificacao)

    def titulo(self, texto):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(limpa_texto(texto).upper())
        r.bold = True
        r.font.size = Pt(self.tam + 1)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        return p

    def identificacao(self, titulo, linha2=None):
        """Bloco de abertura registrável: título centralizado, a qualificação
        da sociedade (razão social, CNPJ, NIRE) e um filete duplo abaixo.

        Os identificadores recebem hífen NÃO SEPARÁVEL: sem isso o Word quebra
        "17.3.0000000-1" ao fim da linha e o NIRE aparece partido em duas — numa
        folha que vai à Junta Comercial isso é defeito, não detalhe."""
        p = self.titulo(titulo)
        p.paragraph_format.space_after = Pt(4)
        if linha2:
            q = self.doc.add_paragraph()
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q.paragraph_format.space_after = Pt(6)
            q.paragraph_format.keep_with_next = True
            r = q.add_run(liga_identificadores(linha2))
            r.font.size = Pt(self.tam - 1)
        fil = self.doc.add_paragraph()
        fil.paragraph_format.space_before = Pt(0)
        fil.paragraph_format.space_after = Pt(14)
        fil.paragraph_format.keep_with_next = True
        pbdr = _el("w:pBdr")
        pbdr.append(_el("w:bottom", val="double", sz=6, space="1", color="000000"))
        fil._p.get_or_add_pPr().append(pbdr)
        return p

    def secao(self, texto, caps=True):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto = limpa_texto(texto)
        r = p.add_run(texto.upper() if caps else texto)
        r.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return p

    def paragrafo(self, texto, recuo=True):
        p = self.doc.add_paragraph(liga_identificadores(texto))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if recuo:
            p.paragraph_format.first_line_indent = Cm(1.25)
        return p

    _JA_NUMERADO = re.compile(r"^\s*(\d+[\.\)]|[IVXLC]+\s*[–\-\.]|§)")

    def item(self, texto, recuo=False):
        """Deliberação/item numerado. O kit numera "1., 2., 3." sozinho — e não
        duplica quando o texto já chega numerado."""
        texto = str(texto)
        if self._JA_NUMERADO.match(texto):
            return self.paragrafo(texto, recuo=recuo)
        self._itens += 1
        return self.paragrafo("%d. %s" % (self._itens, texto), recuo=recuo)

    def clausula(self, titulo, texto=None):
        """"CLÁUSULA PRIMEIRA — DO OBJETO" com o ordinal automático."""
        self._clausulas += 1
        self._incisos = 0
        self._paragrafos_numerados = 0
        ordem = (self.ORDINAIS[self._clausulas - 1] if self._clausulas <= len(self.ORDINAIS)
                 else "%dª" % self._clausulas)
        cabeca = "CLÁUSULA %s — %s" % (ordem, limpa_texto(titulo).upper())
        p = self.secao(cabeca)
        if texto:
            self.paragrafo(texto)
        return p

    def paragrafo_unico(self, texto):
        return self.paragrafo("Parágrafo único. " + limpa_texto(texto))

    def paragrafo_numerado(self, texto):
        """"§ 1º", "§ 2º" — numerados dentro da cláusula corrente."""
        self._paragrafos_numerados += 1
        return self.paragrafo("§ %dº %s" % (self._paragrafos_numerados, limpa_texto(texto)))

    def inciso(self, texto):
        """"I – ", "II – " — numerados dentro da cláusula corrente."""
        self._incisos += 1
        romano = (self.ROMANOS[self._incisos - 1] if self._incisos <= len(self.ROMANOS)
                  else str(self._incisos))
        p = self.paragrafo("%s – %s" % (romano, limpa_texto(texto)), recuo=False)
        p.paragraph_format.left_indent = Cm(1.25)
        return p

    def fecho(self, local_data):
        self._fecho_pendente = str(local_data)
        return self

    def assinaturas(self, nomes, cargos=None, subtitulos=None, local_data=None):
        """Assinaturas em PARES lado a lado — uma folha de ata com seis sócios
        empilhados um por linha vira três páginas de espaço em branco.

        O fecho e as assinaturas ficam numa tabela de uma coluna que NÃO se
        parte: a linha de assinatura nunca sobra sozinha numa página."""
        nomes = list(nomes)
        cargos = list(cargos or subtitulos or []) + [None] * len(nomes)
        if local_data:
            self.fecho(local_data)
        pendente = self._fecho_pendente
        self._fecho_pendente = None
        self._assinantes += [(n, cargos[i]) for i, n in enumerate(nomes)]
        if pendente:
            self._textos_de_fechamento.append(pendente)

        externa = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(externa)
        _nao_quebrar_linha(externa.rows[0])
        celula = externa.rows[0].cells[0]
        _cell_pad(celula, 0, 0, 0, 0)
        _limpa_paragrafo(celula.paragraphs[0])
        if pendente:
            p = celula.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(14)
            p.add_run(limpa_texto(pendente))
        self._pares(celula, nomes, cargos)
        _tblpr_put(externa._tbl.tblPr, _el("w:tblW", w=5000, type="pct"))
        return externa

    def testemunhas(self, nomes, cargos=None):
        """Bloco "TESTEMUNHAS" com as duas assinaturas exigidas em contrato."""
        nomes = list(nomes)
        cargos = list(cargos or []) + [None] * len(nomes)
        externa = self.doc.add_table(rows=1, cols=1)
        self._sem_bordas(externa)
        _nao_quebrar_linha(externa.rows[0])
        celula = externa.rows[0].cells[0]
        _cell_pad(celula, 0, 0, 0, 0)
        _limpa_paragrafo(celula.paragraphs[0])
        p = celula.paragraphs[0]
        p.paragraph_format.space_before = Pt(16)
        r = p.add_run("TESTEMUNHAS:")
        r.bold = True
        r.font.size = Pt(self.tam - 1)
        self._assinantes += [(n, cargos[i]) for i, n in enumerate(nomes)]
        self._pares(celula, nomes, cargos)
        _tblpr_put(externa._tbl.tblPr, _el("w:tblW", w=5000, type="pct"))
        return externa

    def _sem_bordas(self, t):
        b = _el("w:tblBorders")
        for e in ("top", "bottom", "left", "right", "insideH", "insideV"):
            b.append(_el("w:" + e, val="nil"))
        _tblpr_put(t._tbl.tblPr, b)
        return t

    def _pares(self, celula, nomes, cargos):
        # Largura útil desta página: A4 menos as margens oficiais (3 + 2 cm).
        util = 21.0 - 3.0 - 2.0
        meio = 1.2
        w = (util - meio) / 2
        for pi in range(0, len(nomes), 2):
            par = nomes[pi:pi + 2]
            espaco = celula.add_paragraph()
            espaco.paragraph_format.space_after = Cm(1.6 if pi == 0 else 1.1)
            espaco.paragraph_format.line_spacing = Pt(2)
            espaco.add_run().font.size = Pt(2)
            t = celula.add_table(rows=1, cols=3)
            self._sem_bordas(t)
            _nao_quebrar_linha(t.rows[0])
            for j, nome in enumerate(par):
                cell = t.rows[0].cells[j * 2]
                tcPr = cell._tc.get_or_add_tcPr()
                bd = _el("w:tcBorders")
                bd.append(_el("w:top", val="single", sz=6, space="0", color="000000"))
                tcPr.append(bd)
                _cell_pad(cell, 120, 40, 60, 60)
                _limpa_paragrafo(cell.paragraphs[0])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                r = cell.paragraphs[0].add_run(limpa_texto(nome))
                r.bold = True
                r.font.size = Pt(self.tam - 1)
                sub = cargos[pi + j] if pi + j < len(cargos) else None
                if sub:
                    ps = cell.add_paragraph()
                    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ps.paragraph_format.space_after = Pt(0)
                    ps.add_run(liga_identificadores(sub)).font.size = Pt(self.tam - 3)
            # Layout fixo com a grade declarada: sem isso a tabela de assinatura
            # vaza a margem direita da folha registrável.
            larguras = [Cm(x).twips for x in (w, meio, w)]
            _tblpr_put(t._tbl.tblPr, _el("w:tblW", w=sum(larguras), type="dxa"))
            _tblpr_put(t._tbl.tblPr, _el("w:tblLayout", type="fixed"))
            t.autofit = False
            grid = t._tbl.find(qn("w:tblGrid"))
            if grid is not None:
                for col, larg in zip(grid.findall(qn("w:gridCol")), larguras):
                    col.set(qn("w:w"), str(larg))
            for col_i, larg in enumerate(larguras):
                t.rows[0].cells[col_i].width = Twips(larg)

    def _rodape(self):
        f = self.doc.sections[-1].footer
        f.is_linked_to_previous = False
        p = _limpa_paragrafo(f.paragraphs[0])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.rubrica:
            pr = f.add_paragraph()
            pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = pr.add_run("Rubrica: ______________")
            r.font.size = Pt(8)
        p.add_run("Página ").font.size = Pt(9)
        self._campo(p, "PAGE")
        p.add_run(" de ").font.size = Pt(9)
        self._campo(p, "NUMPAGES")

    def _campo(self, p, code):
        r = p.add_run(); r._r.append(_el("w:fldChar", fldCharType="begin"))
        r2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " " + code + " "
        r2._r.append(instr)
        r3 = p.add_run(); r3._r.append(_el("w:fldChar", fldCharType="end"))
        for rr in (r, r2, r3):
            rr.font.size = Pt(9)
            rr.font.color.rgb = RGBColor(0, 0, 0)

    def _metadados(self):
        props = self.doc.core_properties
        props.title = limpa_texto(self.titulo_doc or "Documento")
        props.author = "Frederico AI Studio"
        props.subject = "documento registrável"
        props.language = self.idioma
        props.created = props.modified = datetime.now()

    def salvar(self, caminho, pdf=True, auditar=True):
        """Grava, gera o PDF gêmeo e AUDITA: corpo justificado, zero cor,
        placeholder, identificação quebrada e assinatura órfã."""
        pendente = self._fecho_pendente
        if pendente:
            self._fecho_pendente = None
            p = self.doc.add_paragraph(limpa_texto(pendente))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(14)
        self._metadados()
        self.doc.save(caminho)
        caminho_pdf = _converter_pdf(caminho) if pdf else None
        paginas = _texto_por_pagina(caminho_pdf) if caminho_pdf else []
        if not auditar:
            return relatorio([], paginas=len(paginas), pdf=caminho_pdf)
        rel = relatorio(self._auditar(caminho, caminho_pdf, paginas),
                        paginas=len(paginas), pdf=caminho_pdf)
        return falha_se_grave(rel, "docpro.Sobrio")

    def _auditar(self, caminho, caminho_pdf, paginas):
        from docx import Document as _Doc
        achados = list(self._achados_geracao)
        doc = _Doc(caminho)
        textos = textos_do_docx(doc)
        achados += achados_de_placeholder(textos, "documento registrável")
        # Corpo justificado: o estilo Normal manda, mas um parágrafo pode ter
        # sido alinhado à mão por quem estendeu o script.
        soltos = [p.text[:40] for p in doc.paragraphs
                  if p.text.strip() and p.alignment not in (None, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                                            WD_ALIGN_PARAGRAPH.CENTER)]
        if soltos:
            achados.append(achado("grave", "corpo-nao-justificado",
                                  "%d parágrafo(s) fora do justificado: %s"
                                  % (len(soltos), soltos[:3])))
        coloridos = [r.text[:30] for p in doc.paragraphs for r in p.runs
                     if r.font.color is not None and r.font.color.rgb is not None
                     and str(r.font.color.rgb) not in ("000000",)]
        if coloridos:
            achados.append(achado("grave", "cor-em-documento-sobrio",
                                  "documento registrável não leva cor: %s" % coloridos[:3]))
        quebrados = [t for t in textos if re.search(r"-\s*\n", t)]
        if quebrados:
            achados.append(achado("aviso", "identificador-quebrado",
                                  "identificador possivelmente partido: %s" % quebrados[:2]))
        if caminho_pdf is None:
            achados.append(achado("aviso", "sem-pdf-gemeo",
                                  "o PDF gêmeo não foi gerado (LibreOffice indisponível)"))
            return achados
        achados += kits.achados_de_paginacao(
            paginas, self._assinantes, self._textos_de_fechamento,
            ultima_e_arte=self._contracapa_de_pagina)
        return achados


# atalhos de módulo
def sombrear(cell, cor):
    sombrear_celula(cell, cor)
